from fastapi import FastAPI, WebSocket, WebSocketDisconnect, Depends, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from sqlalchemy.orm import Session
from typing import Dict, Set
import json
import asyncio
from datetime import datetime

from models import *
from database import get_db, init_db
from redis_client import redis_client, get_pubsub
from config import config
from tasks import deep_search_task
from contextlib import asynccontextmanager

@asynccontextmanager
async def lifespan(app: FastAPI):
    # logger.info("🚀 Starting application...")
    
    try:
        Base.metadata.create_all(bind=engine)
    except Exception as e:
        logger.error(f"❌ Database initialization failed: {e}")
    
    try:
        await conversation_manager.connect_redis()
        # logger.info("✅ Redis connected")
    except Exception as e:
        logger.warning(f"⚠️ Redis connection failed: {e}")
     
    yield
    
    # logger.info("🛑 Shutting down...")
    await conversation_manager.disconnect_redis()
    # logger.info("✅ Shutdown complete")

# Initialize FastAPI app
app = FastAPI(title="Deep Search API", lifespan=lifespan)
 
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:8084",      # Your frontend port
        "http://127.0.0.1:8084",
        "http://localhost:8082",
        "http://127.0.0.1:8082",
        "http://localhost:3000",      # Common dev ports
        "http://127.0.0.1:3000",
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["*"]
)

# WebSocket connection manager
class ConnectionManager:
    def __init__(self):
        self.active_connections: Dict[str, Set[WebSocket]] = {}
        self.pubsub_tasks: Dict[str, asyncio.Task] = {}
    
    async def connect(self, websocket: WebSocket, job_id: str):
        if job_id not in self.active_connections:
            self.active_connections[job_id] = set()
        
        self.active_connections[job_id].add(websocket)
        
        if job_id not in self.pubsub_tasks:
            task = asyncio.create_task(self._listen_to_redis(job_id))
            self.pubsub_tasks[job_id] = task
            
    def disconnect(self, websocket: WebSocket, job_id: str):
        if job_id in self.active_connections:
            self.active_connections[job_id].discard(websocket)
            
            if not self.active_connections[job_id]:
                del self.active_connections[job_id]
                
                if job_id in self.pubsub_tasks:
                    self.pubsub_tasks[job_id].cancel()
                    del self.pubsub_tasks[job_id]
    

    async def _listen_to_redis(self, job_id: str):
        pubsub = get_pubsub()
        channel = f"job:{job_id}"
        
        try:
            pubsub.subscribe(channel)
            while True:
                message = pubsub.get_message(ignore_subscribe_messages=True)
                
                if message and message["type"] == "message":
                    data = message["data"]
                    
                    # Send to all connected WebSockets for this job
                    if job_id in self.active_connections:
                        disconnected = set()
                        
                        for websocket in self.active_connections[job_id]:
                            try:
                                await websocket.send_text(data)
                            except Exception as e:
                                print(f"[WS] ❌ Failed to send: {e}")
                                disconnected.add(websocket)
                        
                        # Clean up disconnected WebSockets
                        for ws in disconnected:
                            self.disconnect(ws, job_id)
                        
                        # Check if job is complete
                        try:
                            msg = json.loads(data)
                            if msg.get("type") in ["complete", "error"]:
                                await asyncio.sleep(5)
                                break
                        except json.JSONDecodeError:
                            pass
                
                # Small delay to prevent CPU spinning
                await asyncio.sleep(0.01)
                    
        except asyncio.CancelledError:
            print(f"[WS] ⚠️ Listener cancelled for {job_id}")
        finally:
            pubsub.unsubscribe(channel)
            pubsub.close()
        
manager = ConnectionManager()

@app.on_event("startup")
async def startup_event():
    init_db()

@app.get("/health")
async def health_check():
    return {"status": "healthy", "timestamp": datetime.utcnow().isoformat()}

@app.post("/search")
async def create_search(
    query: str,
    db: Session = Depends(get_db)
):
    if not query or len(query.strip()) == 0:
        raise HTTPException(status_code=400, detail="Query cannot be empty")
    
    # Create job record
    job = SearchJob(query=query, status=JobStatus.PENDING)
    db.add(job)
    db.commit()
    db.refresh(job)
    
    # Dispatch Celery task
    task = deep_search_task.apply_async(
        args=[job.id, query],
        task_id=f"search-{job.id}"
    )
    
    job.celery_task_id = task.id
    db.commit()
    
    return {
        "job_id": job.id,
        "status": job.status.value,
        "message": "Search job created successfully"
    }

@app.get("/search/{job_id}")
async def get_search_status(
    job_id: str,
    db: Session = Depends(get_db)
):
 
    job = db.query(SearchJob).filter(SearchJob.id == job_id).first()
    
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    
    response = job.to_dict()
    
    if job.result:
        try:
            response["result"] = json.loads(job.result)
        except json.JSONDecodeError:
            pass
    
    return response

 
@app.websocket("/ws")
async def websocket_endpoint(websocket: WebSocket):
    await websocket.accept()
    current_job_id = None
    
    try:
        while True:
            data = await websocket.receive_text()
            
            try:
                message = json.loads(data)
                action = message.get("action")
                
                if action == "subscribe":
                    job_id = message.get("job_id")
                    
                    if not job_id:
                        await websocket.send_json({
                            "type": "error",
                            "content": "job_id is required"
                        })
                        continue
                    
                    # Unsubscribe from previous job if any
                    if current_job_id:
                        manager.disconnect(websocket, current_job_id)
                    
                    # Subscribe to new job
                    current_job_id = job_id
                    await manager.connect(websocket, job_id)
                    
                    await websocket.send_json({
                        "type": "subscribed",
                        "content": f"Subscribed to job {job_id}"
                    })
                
                elif action == "unsubscribe":
                    if current_job_id:
                        manager.disconnect(websocket, current_job_id)
                        current_job_id = None
                    
                    await websocket.send_json({
                        "type": "unsubscribed",
                        "content": "Unsubscribed successfully"
                    })
                
                elif action == "ping":
                    await websocket.send_json({
                        "type": "pong",
                        "content": "Connection alive"
                    })
                
                else:
                    await websocket.send_json({
                        "type": "error",
                        "content": f"Unknown action: {action}"
                    })
                    
            except json.JSONDecodeError:
                print(f"[WS_ENDPOINT] ❌ Invalid JSON received")
                await websocket.send_json({
                    "type": "error",
                    "content": "Invalid JSON"
                })
    
    except WebSocketDisconnect:
        print(f"[WS_ENDPOINT] 🔌 WebSocket disconnected")
        if current_job_id:
            print(f"[WS_ENDPOINT] 🧹 Cleaning up job: {current_job_id}")
            
            manager.disconnect(websocket, current_job_id)
from fastapi.staticfiles import StaticFiles
app.mount("/", StaticFiles(directory="static", html=True), name="static")
@app.get("/")
async def root():
    return FileResponse("index.html")
    # return {
    #     "service": "Deep Search API",
    #     "version": "1.0.0",
    #     "endpoints": {
    #         "create_search": "POST /search",
    #         "get_status": "GET /search/{job_id}",
    #         "websocket": "WS /ws",
    #         "health": "GET /health"
    #     }
    # }


from fastapi import FastAPI, HTTPException, WebSocket, WebSocketDisconnect, Depends, UploadFile, File, Header, Request,Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, FileResponse, HTMLResponse, Response
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from pydantic import BaseModel, EmailStr
from sqlalchemy import create_engine, Column, String, DateTime, Text, Boolean, Integer, ForeignKey
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker, Session, relationship
from sqlalchemy.dialects.postgresql import UUID
import uuid
from datetime import datetime, timezone, timedelta
from typing import List, Dict, Optional, Any, AsyncGenerator
import json
import asyncio
import redis.asyncio as aioredis
import logging
from contextlib import asynccontextmanager
import anthropic
import os
from dotenv import load_dotenv
import aiohttp
import hashlib
from jose import jwt
from passlib.context import CryptContext
import base64
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
import time
from collections import defaultdict
import re
# import spacy 
  
# from reportlab.lib.units import inch
from io import BytesIO
import markdown
from bs4 import BeautifulSoup
import markdown
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import httpx

async def generate_message_markdown(message_content: str, sources: list = None) -> str:
    md_content = f"# NOIR AI Response\n\n"
    md_content += f"{message_content}\n\n"
    
    if sources:
        md_content += "## Sources\n\n"
        for i, source in enumerate(sources, 1):
            url = source.get('url') if isinstance(source, dict) else source
            title = source.get('title', '') if isinstance(source, dict) else ''
            if title:
                md_content += f"{i}. [{title}]({url})\n"
            else:
                md_content += f"{i}. {url}\n"
    
    return md_content


async def generate_message_pdf(message_content: str, sources: list = None) -> BytesIO:
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=72, leftMargin=72,
                          topMargin=72, bottomMargin=18)
    
    styles = getSampleStyleSheet()
    story = []
    
    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=12,
        textColor='#000000'
    )
    story.append(Paragraph("NOIR AI Response", title_style))
    story.append(Spacer(1, 0.3*inch))
    
    # Process content line by line
    lines = message_content.split('\n')
    current_para = []
    
    for line in lines:
        line = line.strip()
        
        if not line:
            if current_para:
                story.append(Paragraph(' '.join(current_para), styles['BodyText']))
                story.append(Spacer(1, 0.1*inch))
                current_para = []
            continue
        
        # Handle headings
        if line.startswith('### '):
            if current_para:
                story.append(Paragraph(' '.join(current_para), styles['BodyText']))
                current_para = []
            story.append(Paragraph(line[4:], styles['Heading3']))
            story.append(Spacer(1, 0.1*inch))
        elif line.startswith('## '):
            if current_para:
                story.append(Paragraph(' '.join(current_para), styles['BodyText']))
                current_para = []
            story.append(Paragraph(line[3:], styles['Heading2']))
            story.append(Spacer(1, 0.15*inch))
        elif line.startswith('# '):
            if current_para:
                story.append(Paragraph(' '.join(current_para), styles['BodyText']))
                current_para = []
            story.append(Paragraph(line[2:], styles['Heading1']))
            story.append(Spacer(1, 0.2*inch))
        # Handle code blocks
        elif line.startswith('```'):
            if current_para:
                story.append(Paragraph(' '.join(current_para), styles['BodyText']))
                current_para = []
            continue
        # Handle bullet points
        elif line.startswith('- ') or line.startswith('* '):
            if current_para:
                story.append(Paragraph(' '.join(current_para), styles['BodyText']))
                current_para = []
            story.append(Paragraph(line[2:], styles['Normal'], bulletText='•'))
        else:
            current_para.append(line)
    
    if current_para:
        story.append(Paragraph(' '.join(current_para), styles['BodyText']))
    
    # Add sources
    if sources:
        story.append(Spacer(1, 0.3*inch))
        story.append(Paragraph("Sources", styles['Heading2']))
        story.append(Spacer(1, 0.1*inch))
        
        for i, source in enumerate(sources, 1):
            url = source.get('url') if isinstance(source, dict) else source
            title = source.get('title', '') if isinstance(source, dict) else ''
            source_text = f"{i}. {title}<br/><font size='8'>{url}</font>" if title else f"{i}. {url}"
            story.append(Paragraph(source_text, styles['BodyText']))
            story.append(Spacer(1, 0.08*inch))
    
    doc.build(story)
    buffer.seek(0)
    return buffer

async def generate_message_docx(message_content: str, sources: list = None) -> BytesIO:
    doc = Document()
    
    title = doc.add_heading('NOIR AI Response', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    lines = message_content.split('\n')
    current_para = []
    
    for line in lines:
        line = line.strip()
        
        if not line:
            if current_para:
                doc.add_paragraph(' '.join(current_para))
                current_para = []
            continue
        
        if line.startswith('### '):
            if current_para:
                doc.add_paragraph(' '.join(current_para))
                current_para = []
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            if current_para:
                doc.add_paragraph(' '.join(current_para))
                current_para = []
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            if current_para:
                doc.add_paragraph(' '.join(current_para))
                current_para = []
            doc.add_heading(line[2:], level=1)
        # Handle bullet points
        elif line.startswith('- ') or line.startswith('* '):
            if current_para:
                doc.add_paragraph(' '.join(current_para))
                current_para = []
            doc.add_paragraph(line[2:], style='List Bullet')
        # Handle code blocks
        elif line.startswith('```'):
            continue
        else:
            current_para.append(line)
    
    if current_para:
        doc.add_paragraph(' '.join(current_para))
    
    # Add sources
    if sources:
        doc.add_heading('Sources', level=1)
        for i, source in enumerate(sources, 1):
            url = source.get('url') if isinstance(source, dict) else source
            title = source.get('title', '') if isinstance(source, dict) else ''
            source_text = f"{title} - {url}" if title else url
            doc.add_paragraph(source_text, style='List Number')
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def load_config():
    config_data = {}
    
    if os.path.exists('api_keys.conf'):
        try:
            with open('api_keys.conf', 'r') as f:
                for line in f:
                    line = line.strip()
                    if not line or line.startswith('#'):
                        continue
                    # Parse KEY=VALUE
                    if '=' in line:
                        key, value = line.split('=', 1)
                        config_data[key.strip()] = value.strip()
        except Exception as e:
            logger.error(f"Error reading api_keys.conf: {e}")
    else:
        # Fallback to .env
        # logger.info("api_keys.conf not found, trying .env")
        load_dotenv()
        url = os.getenv('DATABASE_URL')
        
        if url:
            logger.info("DATABASE URL BEING LOADED IS ->" + url)
        else:
            logger.info("DATABASE URL BEING LOADED IS ->" + "None")
            
        config_data['ANTHROPIC_API_KEY'] = os.getenv('CLAUDE_API_KEY', os.getenv('ANTHROPIC_API_KEY', ''))
        config_data['SERPAPI_KEY'] = os.getenv('SERPAPI_KEY',  '')
        config_data['BRIGHT_DATA_KEY'] = os.getenv('BRIGHT_DATA_KEY', '')
        config_data['TWOCAPTCHA_KEY'] = os.getenv('TWOCAPTCHA_KEY', '')
        config_data['ALPHA_VANTAGE_KEY'] = os.getenv('ALPHA_VANTAGE_KEY', 'demo')
        config_data['OPENWEATHER_KEY'] = os.getenv('OPENWEATHER_KEY', 'demo')
        config_data['NEWSAPI_KEY'] = os.getenv('NEWSAPI_KEY', 'demo')
        config_data['DATABASE_URL'] = os.getenv('DATABASE_URL', 'postgresql://chatuser:chatpass123@localhost:5432/chatdb')
        config_data['REDIS_URL'] = os.getenv('REDIS_URL', 'redis://localhost:6379')
        config_data['JWT_SECRET'] = os.getenv('JWT_SECRET', 'your-secret-key-change-in-production')
        config_data['APP_ENV'] = os.getenv('APP_ENV', 'production')
        config_data['LOG_LEVEL'] = os.getenv('LOG_LEVEL', 'info')
        config_data['GOOGLE_API_KEY'] = os.getenv('GOOGLE_API_KEY', '')
        config_data['GOOGLE_CX'] = os.getenv('GOOGLE_CX', '')
        config_data['GOOGLE_CLIENT_ID'] = os.getenv('GOOGLE_CLIENT_ID', '')
        config_data['GOOGLE_CLIENT_SECRET'] = os.getenv('GOOGLE_CLIENT_SECRET', '')
        config_data['REDIRECT_URI'] = os.getenv('REDIRECT_URI', '')
        config_data['SECRET_KEY'] = os.getenv('SECRET_KEY', '')
        config_data['FRONTEND_URL'] = os.getenv('FRONTEND_URL', '')
        config_data['OPENAI_KEY'] = os.getenv('OPENAI_KEY', '')
        config_data['SMTP_USERNAME'] = os.getenv('SMTP_USERNAME', '')
        config_data['SMTP_PASSWORD'] = os.getenv('SMTP_PASSWORD', '')
        config_data['SMTP_HOST'] = os.getenv('SMTP_HOST', '')
        config_data['SMTP_PORT'] = os.getenv('SMTP_PORT', '')
        config_data['CLAUDE_API_KEY'] = os.getenv('CLAUDE_API_KEY', '')
 
    return config_data

# Load configuration
config = load_config()
DATABASE_URL = config.get('DATABASE_URL', 'postgresql://chatuser:chatpass123@localhost:5432/chatdb')
REDIS_URL = config.get('REDIS_URL', 'redis://localhost:6379')
CLAUDE_API_KEY = config.get('ANTHROPIC_API_KEY', '')
SERPAPI_KEY = config.get('SERPAPI_KEY', '')
TWOCAPTCHA_KEY = config.get('TWOCAPTCHA_KEY', '')
ALPHA_VANTAGE_KEY = config.get('ALPHA_VANTAGE_KEY', 'demo')
OPENWEATHER_KEY = config.get('OPENWEATHER_KEY', 'demo')
NEWSAPI_KEY = config.get('NEWSAPI_KEY', 'demo')
JWT_SECRET = config.get('JWT_SECRET', 'your-secret-key-change-in-production')
APP_ENV = config.get('APP_ENV', 'production')
LOG_LEVEL = config.get('LOG_LEVEL', 'info').upper()
JWT_ALGORITHM = "HS256"
OPENAI_KEY =  config.get('OPENAI_KEY', '')
BRIGHT_DATA_KEY = config.get('BRIGHT_DATA_KEY', '') 
GOOGLE_API_KEY =  config.get('GOOGLE_API_KEY', '')
GOOGLE_CX = config.get('GOOGLE_CX', '') 
GOOGLE_CLIENT_ID = config.get('GOOGLE_CLIENT_ID', '')
GOOGLE_CLIENT_SECRET = config.get('GOOGLE_CLIENT_SECRET', '') 
REDIRECT_URI =  config.get('REDIRECT_URI', '')
SECRET_KEY = config.get('SECRET_KEY', '') 
FRONTEND_URL = config.get('FRONTEND_URL', '')  
SMTP_USERNAME = config.get('SMTP_USERNAME', '') 
SMTP_PASSWORD =  config.get('SMTP_PASSWORD', '')
SMTP_HOST = config.get('SMTP_HOST', '') 
SMTP_PORT = config.get('FRONTEND_URL', '')  
CLAUDE_API_KEY = config.get('CLAUDE_API_KEY', '')  

logging.getLogger().setLevel(getattr(logging, LOG_LEVEL))

pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

# class UserCreate(BaseModel):
#     username: str
#     email: str
#     password: str

# class UserLogin(BaseModel):
#     email: str
#     password: str

# class ConversationCreate(BaseModel):
#     title: Optional[str] = "New Conversation"

# class MessageCreate(BaseModel):
#     content: str

# class MessageSend(BaseModel):
#     content: str
#     conversation_id: Optional[str] = None
#     deep_search: Optional[bool] = False    
#     lab_mode: Optional[bool] = False  # NEW

# class ReactionCreate(BaseModel):
#     reaction_type: str

# class MessageResponse(BaseModel):
#     id: str
#     role: str
#     content: str
#     has_file: bool
#     created_at: datetime
#     reactions: List[Dict]

# class ConversationResponse(BaseModel):
#     id: str
#     title: str
#     created_at: datetime
#     updated_at: datetime
#     message_count: int
 
class ConversationManager: 

    def __init__(self, redis_url: str = REDIS_URL):
        self.claude_client = anthropic.Anthropic(api_key=CLAUDE_API_KEY)
        
        self.redis_url = redis_url
        self.redis = None
        self.http_session = None
        self.serpapi_key = SERPAPI_KEY
        # self.serpapi_key = None
        self.openweather_key = OPENWEATHER_KEY
        self.newsapi_key = NEWSAPI_KEY
        # self.nlp = nlp  # spaCy model 
        self.bright_data_key = BRIGHT_DATA_KEY 
        # self.rate_limit_handler = RateLimitHandler() 
        self.google_api_key = GOOGLE_API_KEY
        self.google_cx = GOOGLE_CX
     
    
    async def connect_redis(self):
        if not self.redis:
            self.redis = await aioredis.from_url(self.redis_url, decode_responses=True)
        
        if not self.http_session:
            timeout = aiohttp.ClientTimeout(total=5, connect=2)
            connector = aiohttp.TCPConnector(
                limit=100,
                limit_per_host=10,
                ttl_dns_cache=300
            )
            self.http_session = aiohttp.ClientSession(
                timeout=timeout,
                connector=connector,
                headers={
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
                }
            )

    async def disconnect_redis(self):
        if self.http_session:
            await self.http_session.close()
            # logger.info("HTTP connection pool closed")
        
        if self.redis:
            await self.redis.close()
            # logger.info("Disconnected from Redis")

    async def get_conversation_history(self, conversation_id: str, db: Session) -> List[Dict]:
        cache_key = f"conv:{conversation_id}:history"
        cached = await self.redis.get(cache_key)
        if cached:
            # logger.info(f"Cache hit for conversation {conversation_id}")
            return json.loads(cached)
        
        messages = db.query(Message).filter(
            Message.conversation_id == uuid.UUID(conversation_id)
        ).order_by(Message.created_at).all()
        
        history = []
        for m in messages:
            msg_dict = {"role": m.role, "content": m.content}
            if m.has_file and m.file_data:
                msg_dict["content"] = [
                    {
                        "type": "image",
                        "source": {
                            "type": "base64",
                            "media_type": m.file_type or "image/jpeg",
                            "data": m.file_data
                        }
                    },
                    {"type": "text", "text": m.content}
                ]
            history.append(msg_dict)
        
        await self.redis.setex(cache_key, 300, json.dumps(history, default=str))
        return history

    async def scrape_url(self, url: str, timeout: int = 5) -> Optional[str]:
        try:
            cache_key = f"scraped:{hashlib.md5(url.encode()).hexdigest()}"
            cached = await self.redis.get(cache_key)
            if cached:
                return cached
            
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
                'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
                'Accept-Language': 'en-US,en;q=0.5'
            }
            
            async with self.http_session.get(url, headers=headers) as response:
                if response.status != 200:
                    logger.warning(f"[SCRAPER] HTTP {response.status} for {url}")
                    return None
                
                html = await response.text()
                
                from bs4 import BeautifulSoup
                soup = BeautifulSoup(html, 'html.parser')
                
                for script in soup(["script", "style", "nav", "footer", "header", "aside", "iframe", "noscript"]):
                    script.decompose()
                
                main_content = None
                for selector in ['main', 'article', '[role="main"]', '.content', '#content']:
                    main_content = soup.select_one(selector)
                    if main_content:
                        break
                
                if not main_content:
                    main_content = soup.find('body')
                
                if not main_content:
                    return None
                
                text = main_content.get_text(separator=' ', strip=True)
                text = re.sub(r'\s+', ' ', text).strip()
                
                if len(text) > 1000:
                    text = text[:1000] + "..."
                
                if text:
                    await self.redis.setex(cache_key, 3600, text)
                
                return text
                
        except asyncio.TimeoutError:
            logger.warning(f"[SCRAPER] ⏱️ Timeout scraping {url}")
            return None
        except Exception as e:
            logger.error(f"[SCRAPER] ❌ Error scraping {url}: {str(e)}")
            return None
        
conversation_manager = ConversationManager()
 
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

# Auth helpers
def create_access_token(data: dict):
    to_encode = data.copy()
    expire = datetime.now(timezone.utc) + timedelta(hours=24)
    to_encode.update({"exp": expire})
    return jwt.encode(to_encode, JWT_SECRET, algorithm=JWT_ALGORITHM)

def verify_token(authorization: str = Header(None)) -> dict:
    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Missing or invalid token")
    
    token = authorization.replace("Bearer ", "")
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        return payload
    except Exception as e:
        print("Invalid token -> " + str(e))
        raise HTTPException(status_code=401, detail=str(e))

def get_current_user(authorization: str = Header(None)) -> Optional[dict]:
    if not authorization:
        return None
    try:
        return verify_token(authorization)
    except:
        return None

def get_current_user_optional(authorization: str = Header(None)) -> Optional[dict]:
    if not authorization:
        return None
    
    try:
        # Extract token from "Bearer <token>" format
        if authorization.startswith("Bearer "):
            token = authorization.replace("Bearer ", "")
        else:
            token = authorization
        
        # Decode JWT token using your existing settings
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        user_id = payload.get("user_id")
        username = payload.get("username")
        
        if not user_id:
            return None
        
        return {
            "user_id": user_id,
            "username": username,
            "email": payload.get("email")
        }
    except jwt.ExpiredSignatureError:
        # Token expired - return None (allow anonymous access)
        return None
    except jwt.JWTError:
        # Invalid token - return None (allow anonymous access)
        return None
    except Exception as e:
        print(f"Auth error: {e}")
        return None
 
@app.post("/chat/send")
async def send_chat_message(
    message: MessageSend,
    db: Session = Depends(get_db),
    current_user: Optional[dict] = Depends(get_current_user)
):
    """Send message - auto-creates conversation if needed. Works with/without auth."""
    conversation_id = message.conversation_id
    
    if not conversation_id:
        conversation = Conversation(
            user_id=uuid.UUID(current_user["user_id"]) if current_user else None,
            title="New Conversation",
            is_anonymous=current_user is None
        )
        db.add(conversation)
        db.commit()
        db.refresh(conversation)
        conversation_id = str(conversation.id)
        # # logger.info(f"Auto-created conversation: {conversation_id}")
    
    conv = db.query(Conversation).filter(Conversation.id == uuid.UUID(conversation_id)).first()
    
    if not conv:
        raise HTTPException(status_code=404, detail="Conversation not found")
    
    if conv.is_anonymous and conv.message_count >= 10:
        return {
            "role": "assistant",
            "content": "You've reached the free message limit (10 messages). Please register for unlimited access!",
            "limit_reached": True,
            "conversation_id": conversation_id,
            "message_count": conv.message_count
        }
    
    response = await conversation_manager.chat_with_claude_enhanced(conversation_id, message.content, db)
    
    return {
        "role": "assistant",
        "content": response,
        "conversation_id": conversation_id,
        "message_count": conv.message_count,
        "is_anonymous": conv.is_anonymous
    }

@app.options("/chat/send/stream")
async def chat_stream_options(request: Request):
    """Handle preflight CORS request"""
    return Response(
        status_code=200,
        headers={
            "Access-Control-Allow-Origin": request.headers.get("origin", "*"),
            "Access-Control-Allow-Credentials": "true",
            "Access-Control-Allow-Methods": "POST, OPTIONS",
            "Access-Control-Allow-Headers": "Content-Type, Authorization, *",
            "Access-Control-Max-Age": "86400"
        }
    )


from datetime import timedelta

async def cleanup_old_anonymous_usage(db: Session, days_to_keep: int = 7):
    """
    Remove anonymous usage records older than X days.
    """
    cutoff_date = date.today() - timedelta(days=days_to_keep)
    
    db.query(AnonymousUsage).filter(
        AnonymousUsage.usage_date < cutoff_date
    ).delete()
    
    db.commit()
    
@app.post("/chat/send/stream") 
async def send_chat_message_stream(
    request: Request,
    content: str = Form(...),
    conversation_id: Optional[str] = Form(None),
    deep_search: Optional[bool] = Form(False),
    lab_mode: Optional[bool] = Form(False),
    files: List[UploadFile] = File(default=[]),
    db: Session = Depends(get_db),
    current_user: Optional[dict] = Depends(get_current_user)
):
    message = MessageSend(
        content=content,
        conversation_id=conversation_id,
        deep_search=deep_search,
        lab_mode=lab_mode
    )
    
    client_ip = get_client_ip(request)
    
    # Check if anonymous user
    is_anonymous = current_user is None
    
    if is_anonymous:
        # Check daily limit by IP
        is_allowed, remaining = await check_anonymous_limit(db, client_ip)
        
        if not is_allowed:
            async def limit_response():
                yield json.dumps({
                    "type": "error",
                    "message": "Daily message limit reached for anonymous users",
                    "limit_reached": True,
                    "remaining": 0
                }) + "\n"
                yield json.dumps({"type": "done"}) + "\n"
            
            return StreamingResponse(
                limit_response(),
                media_type="text/plain"
            )
        
        # Increment usage counter
        await increment_anonymous_usage(db, client_ip)
    
    
    else:
        user_id = current_user.get("user_id")
        user = db.query(User).filter(User.id == user_id).first()
        if current_user:
            user_id = current_user.get("user_id")
            user = db.query(User).filter(User.id == user_id).first()
            
            # If user has premium credits, deduct one
            if user.message_credits > 0:
                user.message_credits -= 1
                db.commit()
        else:
            # Check daily limit for logged-in user
            is_allowed, remaining = await check_user_daily_limit(db, user_id)
            
            if not is_allowed:
                async def limit_response():
                    yield json.dumps({
                        "type": "error",
                        "message": f"Daily message limit reached ({LOGGED_IN_DAILY_LIMIT} messages per day). Upgrade for unlimited messages!",
                        "limit_reached": True,
                        "user_limit": True,  # ✅ Indicate this is a user limit
                        "remaining": 0
                    }) + "\n"
                    yield json.dumps({"type": "done"}) + "\n"
                
                return StreamingResponse(
                    limit_response(),
                    media_type="text/plain"
                )
            
            # Increment user's daily usage
            await increment_user_daily_usage(db, user_id)
        
    conversation_id = message.conversation_id
    newConversation = False
    
    try:
        # ===== CREATE OR GET CONVERSATION =====
        if not conversation_id:
            conversation = Conversation(
                user_id=uuid.UUID(current_user["user_id"]) if current_user else None,
                title="New Conversation",
                is_anonymous=current_user is None
            )
            newConversation = True
            db.add(conversation)
            db.commit()
            db.refresh(conversation)
            conversation_id = str(conversation.id)
            conv = conversation
        else:
            conv = db.query(Conversation).filter(
                Conversation.id == uuid.UUID(conversation_id)
            ).with_for_update().first()
            
            if not conv:
                raise HTTPException(status_code=404, detail="Conversation not found")
        
    except Exception as e:
        print(f"❌ EXCEPTION WHEN CREATING/GETTING CONV: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))
    
    # ===== CHECK ANONYMOUS MESSAGE LIMIT =====
    # if conv.is_anonymous and conv.message_count >= 2:
    #     async def limit_response():
    #         yield json.dumps({
    #             "type": "error",
    #             "message": "Message limit reached",
    #             "limit_reached": True
    #         }) + "\n"
    #         yield json.dumps({"type": "done"}) + "\n"
        
    #     return StreamingResponse(
    #         limit_response(),
    #         media_type="text/plain",
    #         headers={
    #             "Cache-Control": "no-cache",
    #             "X-Accel-Buffering": "no",
    #             "Access-Control-Allow-Origin": "*",
    #             "Access-Control-Allow-Methods": "POST, OPTIONS",
    #             "Access-Control-Allow-Headers": "*"
    #         }
    #     )
    
    # ===== ROUTE TO CELERY OR DIRECT =====
    use_celery = message.deep_search or message.lab_mode
    
    if use_celery:
        return StreamingResponse(
            stream_response_celery(newConversation, message, conversation_id, conv, files, db),
            media_type="text/plain",
            headers={
                "Cache-Control": "no-cache",
                "X-Accel-Buffering": "no",
                "Connection": "keep-alive",
                "Access-Control-Allow-Origin": "*",
                "Access-Control-Allow-Methods": "POST, OPTIONS",
                "Access-Control-Allow-Headers": "*"
            }
        )
    else:
        return StreamingResponse(
            stream_response_direct(newConversation, message, conversation_id, conv, files, db),
            media_type="text/plain",
            headers={
                "Cache-Control": "no-cache",
                "X-Accel-Buffering": "no",
                "Connection": "keep-alive",
                "Access-Control-Allow-Origin": "*",
                "Access-Control-Allow-Methods": "POST, OPTIONS",
                "Access-Control-Allow-Headers": "*"
            }
        )

async def stream_response_direct(
    newConversation: bool,
    message: MessageSend,
    conversation_id: str,
    conv: Conversation,
    uploaded_files: List[UploadFile],
    db: Session
):
    
    finalSources = []
    extracted_assets = []
    file_contents = [] 
    search_content = ""
    full_response = ""
    app = None
    transformed_query = ""
 
    try:
        is_lab_mode = message.lab_mode
        is_deep_search = message.deep_search
        
        if uploaded_files and len(uploaded_files) > 0:
            
            yield json.dumps({
                "type": "reasoning",
                "step": "File Processing",
                "content": f"Processing {len(uploaded_files)} uploaded file(s)...",
                "timestamp": datetime.now(timezone.utc).isoformat()
            }) + "\n"
            await asyncio.sleep(0.2)
            
            for idx, file in enumerate(uploaded_files, 1):
                try:
                    content_bytes = await file.read()
                    file_size_kb = len(content_bytes) / 1024
                    
                    if file.content_type == "application/pdf":
                        import pdfplumber
                        from io import BytesIO
                        
                        text_content = ""
                        try:
                            with pdfplumber.open(BytesIO(content_bytes)) as pdf:
                                for page_num, page in enumerate(pdf.pages, 1):
                                    page_text = page.extract_text()
                                    if page_text:
                                        text_content += f"\n--- Page {page_num} ---\n{page_text}\n"
                            
                            file_contents.append({
                                "filename": file.filename,
                                "type": "pdf",
                                "pages": len(pdf.pages),
                                "content": text_content[:10000]
                            })
                        except Exception as e:
                            logger.error(f"[FILES] PDF extraction error: {e}")
                            file_contents.append({
                                "filename": file.filename,
                                "type": "pdf",
                                "content": f"Error extracting PDF: {str(e)}"
                            })
                    
                    elif file.content_type in ["text/csv", "application/vnd.ms-excel", 
                                            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"]:
                        import pandas as pd
                        from io import BytesIO
                        
                        try:
                            if file.content_type == "text/csv":
                                df = pd.read_csv(BytesIO(content_bytes))
                            else:
                                df = pd.read_excel(BytesIO(content_bytes))
                            
                            summary = f"File: {file.filename}\n"
                            summary += f"Rows: {len(df)}, Columns: {len(df.columns)}\n"
                            summary += f"Column Names: {', '.join(df.columns.tolist())}\n\n"
                            summary += f"First 10 rows:\n{df.head(10).to_string()}\n\n"
                            
                            if len(df) > 10:
                                summary += f"Summary Statistics:\n{df.describe().to_string()}"
                            
                            file_contents.append({
                                "filename": file.filename,
                                "type": "tabular",
                                "rows": len(df),
                                "columns": len(df.columns),
                                "content": summary[:10000]
                            })
                            
                        except Exception as e:
                            logger.error(f"[FILES] Pandas error: {e}")
                            file_contents.append({
                                "filename": file.filename,
                                "type": "unknown",
                                "content": f"Error processing file: {str(e)}"
                            })
                    
                    elif file.content_type == "text/plain":
                        try:
                            text_content = content_bytes.decode('utf-8', errors='ignore')
                            file_contents.append({
                                "filename": file.filename,
                                "type": "text",
                                "content": text_content[:10000]
                            })
                            logger.info(f"[FILES] ✓ Extracted {len(text_content)} chars from text file")
                        except Exception as e:
                            logger.error(f"[FILES] Text file error: {e}")
                            file_contents.append({
                                "filename": file.filename,
                                "type": "text",
                                "content": f"Error reading text file: {str(e)}"
                            })
                    
                    else:
                        logger.warning(f"[FILES] Unsupported file type: {file.content_type}")
                        file_contents.append({
                            "filename": file.filename,
                            "type": "unsupported",
                            "content": f"Unsupported file type: {file.content_type}"
                        })
                        continue
                    
                    # Send success notification
                    yield json.dumps({
                        "type": "reasoning",
                        "step": f"File {idx} Processed",
                        "content": f"✓ Extracted content from {file.filename}",
                        "timestamp": datetime.now(timezone.utc).isoformat()
                    }) + "\n"
                    await asyncio.sleep(0.1)
                    
                except Exception as e:
                    logger.error(f"[FILES] Error processing {file.filename}: {e}", exc_info=True)
                    yield json.dumps({
                        "type": "reasoning",
                        "step": f"File {idx} Error",
                        "content": f"Failed to process {file.filename}: {str(e)}",
                        "timestamp": datetime.now(timezone.utc).isoformat()
                    }) + "\n"
        
        metadata = json.dumps({
            "type": "metadata",
            "conversation_id": conversation_id,
            "is_anonymous": conv.is_anonymous,
            "deep_search": is_deep_search,
            "lab_mode": is_lab_mode,
            "files_processed": len(file_contents)
        }) + "\n"
        yield metadata
        
        db_messages = db.query(Message).filter(
            Message.conversation_id == uuid.UUID(conversation_id)
        ).order_by(Message.created_at).all()
        
        messages_list = [{"role": msg.role, "content": msg.content} for msg in db_messages]
        
        reasoning_steps = []
        
        from simple_search import simple_search_chat_agent
        user_prompt = message.content
        
        if file_contents:
            user_prompt += "\n\n=== UPLOADED FILES ===\n"
            for file_info in file_contents:
                user_prompt += f"\n--- {file_info['filename']} ({file_info['type']}) ---\n"
                user_prompt += file_info['content']
                user_prompt += "\n" + "="*50 + "\n"
            user_prompt += "\nPlease use the uploaded files content above in answering the prompt.\n\n"
        
        async for results in simple_search_chat_agent(user_prompt, messages_list):
            try:
                data = json.loads(results)
                
                # Handle transformed query
                if data.get("type") == "transformed_query":
                    text = data.get("query", "")                             
                    transformed_query = text
                
                # Handle content streaming
                if data.get("type") == "content":
                    text = data.get("text", "")
                    yield json.dumps({"type": "content", "text": text}) + "\n"
                    full_response += text
                
                # Handle sources
                if data.get("type") == "sources":
                    urls = data.get("urls", [])
                    
                    step = {
                        "type": "reasoning",
                        "step": "Sources Found",
                        "content": transformed_query,
                        "found_sources": len(urls),
                        "sources": urls,
                        "query": transformed_query,
                        "category": "Web Search",
                        "timestamp": datetime.now(timezone.utc).isoformat()
                    }
                    yield json.dumps(step) + "\n" 
                    reasoning_steps.append(step)
                    finalSources.append(urls)

            except Exception as e:
                print(f"Exception while yielding -> {e}")
        user_msg = Message(
            conversation_id=uuid.UUID(conversation_id),
            role="user",
            content=message.content,
            sources=None,
            reasoning_steps=None,
            assets=None,
            lab_mode=True,
            has_file=len(file_contents) > 0,
            file_type=", ".join([f["type"] for f in file_contents]) if file_contents else None
        )
        db.add(user_msg) 
        print(f"SAVED USER MSG, CONV ID -> {conversation_id}") 
        
        assistant_msg = Message(
            conversation_id=uuid.UUID(conversation_id),
            role="assistant",
            content=full_response,
            sources=json.dumps(finalSources) if finalSources else None,
            reasoning_steps=json.dumps(reasoning_steps) if reasoning_steps else None,
            assets=json.dumps(extracted_assets) if extracted_assets else None,
            lab_mode=is_lab_mode,
            app="app",
            celery_task_id=None  # ✅ NULL for direct responses
        )
        db.add(assistant_msg)
        
        conv.updated_at = datetime.now(timezone.utc)
        conv.message_count = (conv.message_count or 0) + 2
        if conv.title == "New Conversation":
            conv.title = message.content[:50] + ("..." if len(message.content) > 50 else "")
        
        db.commit() 
        
        if conversation_manager.redis:
            await conversation_manager.redis.delete(f"conv:{conversation_id}:history")
        
        yield json.dumps({"type": "done"}) + "\n"
        
    except Exception as e:
        logger.error(f"Stream error: {e}", exc_info=True)
        yield json.dumps({
            "type": "error",
            "message": "An error occurred while processing your request"
        }) + "\n"
        yield json.dumps({"type": "done"}) + "\n"

async def process_uploaded_files(uploaded_files: List[UploadFile]) -> list:
    file_contents = []
    
    for file in uploaded_files:
        content_bytes = await file.read()
        
        if file.content_type == "application/pdf":
            import pdfplumber
            from io import BytesIO
            
            text_content = ""
            with pdfplumber.open(BytesIO(content_bytes)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content += page_text + "\n"
            
            file_contents.append({
                "filename": file.filename,
                "type": "pdf",
                "content": text_content[:10000]
            })
        
        elif file.content_type in ["text/csv", "application/vnd.ms-excel", 
                                   "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"]:
            import pandas as pd
            from io import BytesIO
            
            if file.content_type == "text/csv":
                df = pd.read_csv(BytesIO(content_bytes))
            else:
                df = pd.read_excel(BytesIO(content_bytes))
            
            summary = f"File: {file.filename}\n"
            summary += f"Rows: {len(df)}, Columns: {len(df.columns)}\n"
            summary += f"Data:\n{df.head(10).to_string()}"
            
            file_contents.append({
                "filename": file.filename,
                "type": "tabular",
                "content": summary[:10000]
            })
        
        elif file.content_type == "text/plain":
            text_content = content_bytes.decode('utf-8', errors='ignore')
            file_contents.append({
                "filename": file.filename,
                "type": "text",
                "content": text_content[:10000]
            })
    
    return file_contents
 
async def stream_response_celery(
    newConversation: bool,
    message: MessageSend,
    conversation_id: str,
    conv: Conversation,
    uploaded_files: List[UploadFile],
    db: Session
):
    assistant_msg_id = None
    reasoning_step_counter = 0
    
    try:
        # Process uploaded files
        file_contents = []
        if uploaded_files and len(uploaded_files) > 0:
            yield json.dumps({
                "type": "reasoning",
                "step": "File Processing",
                "content": f"Processing {len(uploaded_files)} uploaded file(s)...",
                "timestamp": datetime.now(timezone.utc).isoformat()
            }) + "\n"
            
            file_contents = await process_uploaded_files(uploaded_files)
        
        # Send metadata
        yield json.dumps({
            "type": "metadata",
            "conversation_id": conversation_id,
            "is_anonymous": conv.is_anonymous,
            "deep_search": message.deep_search,
            "lab_mode": message.lab_mode,
            "files_processed": len(file_contents)
        }) + "\n"
        
        # 1. SAVE USER MESSAGE IMMEDIATELY (complete, not placeholder)
        user_msg = Message(
            conversation_id=uuid.UUID(conversation_id),
            role="user",
            content=message.content,
            has_file=len(file_contents) > 0,
            file_type=", ".join([f["type"] for f in file_contents]) if file_contents else None,
            created_at=datetime.now(timezone.utc)
        )
        db.add(user_msg)
        db.commit()
        db.refresh(user_msg)
        print(f"✅ Saved user message: {user_msg.id}")
        
        # Generate job ID and dispatch Celery task
        job_id = str(uuid.uuid4())
        
        # Get conversation history
        db_messages = db.query(Message).filter(
            Message.conversation_id == uuid.UUID(conversation_id)
        ).order_by(Message.created_at).all()
        messages_list = [{"role": msg.role, "content": msg.content} for msg in db_messages]
        
        # Dispatch task
        from tasks import deep_search_task
        task = deep_search_task.apply_async(
            args=[
                job_id,
                conversation_id,
                message.content,
                messages_list,
                file_contents,
                message.lab_mode
            ],
            task_id=f"search-{job_id}",
            queue='llm_worker_queue' 
        )
        
        from redis_client import get_pubsub
        pubsub = get_pubsub()
        channel = f"job:{job_id}"
        pubsub.subscribe(channel)
        
        timeout_seconds = 30 * 60  # 30 minutes
        start_time = time.time()
        final_result = None
        
        while True:
            # Check timeout
            if time.time() - start_time > timeout_seconds:
                from celery_app import celery_app
                celery_app.control.revoke(task.id, terminate=True)
                
                if assistant_msg_id:
                    assistant_msg = db.query(Message).filter(
                        Message.id == assistant_msg_id
                    ).first()
                    if assistant_msg:
                        assistant_msg.status = "failed"
                        assistant_msg.content = "We're sorry, but your request took too long to process. Please try again with a simpler query."
                        db.commit()
                else:
                    error_msg = Message(
                        conversation_id=uuid.UUID(conversation_id),
                        role="assistant",
                        content="We're sorry, but your request took too long to process. Please try again with a simpler query.",
                        celery_task_id=task.id,
                        status="failed",
                        created_at=datetime.now(timezone.utc)
                    )
                    db.add(error_msg)
                    db.commit()
                
                # Update conversation
                conv.updated_at = datetime.now(timezone.utc)
                conv.message_count = (conv.message_count or 0) + 2
                if conv.title == "New Conversation":
                    conv.title = message.content[:50] + ("..." if len(message.content) > 50 else "")
                db.commit()
                
                # Yield error
                yield json.dumps({
                    "type": "error",
                    "message": "Request timeout - please try again"
                }) + "\n"
                
                yield json.dumps({"type": "done"}) + "\n"
                break
            
            # Get message from Redis
            msg = pubsub.get_message(ignore_subscribe_messages=True)
            
            if msg and msg["type"] == "message":
                data_str = msg["data"]
                data = json.loads(data_str)
                
                msg_type = data.get("type")
                
                if msg_type == "reasoning":
                    # Yield reasoning to client (streaming)
                    yield data_str + "\n"
                    
                    # 2. CREATE PLACEHOLDER ASSISTANT MESSAGE on first reasoning step
                    if assistant_msg_id is None:
                        assistant_msg = Message(
                            conversation_id=uuid.UUID(conversation_id),
                            role="assistant",
                            content="",  # Placeholder - will be updated
                            celery_task_id=task.id,
                            status="streaming",
                            created_at=datetime.now(timezone.utc)
                        )
                        db.add(assistant_msg)
                        db.commit()
                        db.refresh(assistant_msg)
                        assistant_msg_id = assistant_msg.id
                        print(f"✅ Created assistant message placeholder: {assistant_msg_id}")
                    
                    try:
                        reasoning_step = ReasoningStep(
                            message_id=assistant_msg_id,
                            step_number=reasoning_step_counter,
                            content=data.get("content", ""),
                            query=data.get("query"),
                            category=data.get("category"),
                            sources=json.dumps(data.get("sources")) if data.get("sources") else None,
                            created_at=datetime.now(timezone.utc)
                        )
                        db.add(reasoning_step)
                        db.commit()
                        reasoning_step_counter += 1
                    except Exception as e:
                        print(f"⚠️ Could not save reasoning step: {e}")
                        # Continue even if reasoning step fails
                
                elif msg_type == "content":
                    yield data_str + "\n"
                    
                    if assistant_msg_id is None:
                        assistant_msg = Message(
                            conversation_id=uuid.UUID(conversation_id),
                            role="assistant",
                            content=data.get("text", ""),
                            celery_task_id=task.id,
                            status="streaming",
                            created_at=datetime.now(timezone.utc)
                        )
                        db.add(assistant_msg)
                        db.commit()
                        db.refresh(assistant_msg)
                        assistant_msg_id = assistant_msg.id
                        print(f"✅ Created assistant message placeholder: {assistant_msg_id}")
                    else:
                        assistant_msg = db.query(Message).filter(
                            Message.id == assistant_msg_id
                        ).first()
                        if assistant_msg:
                            assistant_msg.content += data.get("text", "")
                            db.commit()
                
                elif msg_type == "complete":
                    final_result = data
                    print(f"✅ Received complete result from worker")
                    
                    if assistant_msg_id:
                        assistant_msg = db.query(Message).filter(
                            Message.id == assistant_msg_id
                        ).first()
                        
                        if assistant_msg:
                            assistant_msg.content = final_result.get("content", "")
                            assistant_msg.sources = json.dumps(final_result.get("sources")) if final_result.get("sources") else None
                            # Don't save reasoning_steps here - they're already in ReasoningStep table
                            assistant_msg.assets = json.dumps(final_result.get("assets")) if final_result.get("assets") else None
                            assistant_msg.app = final_result.get("app")
                            assistant_msg.lab_mode = final_result.get("lab_mode", False)
                            assistant_msg.status = "complete"
                            db.commit()
                    else:
                        assistant_msg = Message(
                            conversation_id=uuid.UUID(conversation_id),
                            role="assistant",
                            content=final_result.get("content", ""),
                            sources=json.dumps(final_result.get("sources")) if final_result.get("sources") else None,
                            reasoning_steps=json.dumps(final_result.get("reasoning_steps")) if final_result.get("reasoning_steps") else None,
                            assets=json.dumps(final_result.get("assets")) if final_result.get("assets") else None,
                            app=final_result.get("app"),
                            lab_mode=final_result.get("lab_mode", False),
                            celery_task_id=task.id,
                            status="complete",
                            created_at=datetime.now(timezone.utc)
                        )
                        db.add(assistant_msg)
                        db.commit()
                    
                    conv.updated_at = datetime.now(timezone.utc)
                    conv.message_count = (conv.message_count or 0) + 2
                    if conv.title == "New Conversation":
                        conv.title = message.content[:50] + ("..." if len(message.content) > 50 else "")
                    
                    db.commit()
                    
                    # Clear cache
                    if conversation_manager.redis:
                        await conversation_manager.redis.delete(f"conv:{conversation_id}:history")
                    
                    # Yield done signal
                    yield json.dumps({"type": "done"}) + "\n"
                    break
                
                elif msg_type == "error":
                    # Handle error
                    error_message = data.get("message", "An error occurred")
                    print(f"❌ Worker error: {error_message}")
                    
                    if assistant_msg_id:
                        # Update existing placeholder
                        assistant_msg = db.query(Message).filter(
                            Message.id == assistant_msg_id
                        ).first()
                        if assistant_msg:
                            assistant_msg.status = "failed"
                            assistant_msg.content = f"Error: {error_message}"
                            db.commit()
                    else:
                        # Create error message
                        error_msg = Message(
                            conversation_id=uuid.UUID(conversation_id),
                            role="assistant",
                            content=f"Error: {error_message}",
                            celery_task_id=task.id,
                            status="failed",
                            created_at=datetime.now(timezone.utc)
                        )
                        db.add(error_msg)
                        db.commit()
                    
                    # Update conversation
                    conv.updated_at = datetime.now(timezone.utc)
                    conv.message_count = (conv.message_count or 0) + 2
                    if conv.title == "New Conversation":
                        conv.title = message.content[:50] + ("..." if len(message.content) > 50 else "")
                    
                    db.commit()
                    
                    # Yield error
                    yield data_str + "\n"
                    yield json.dumps({"type": "done"}) + "\n"
                    break
            
            # Small delay to prevent CPU spinning
            await asyncio.sleep(0.01)
        
        # Cleanup
        pubsub.unsubscribe(channel)
        pubsub.close()
        
    except Exception as e:
        logger.error(f"Celery stream error: {e}", exc_info=True)
        
        # Mark message as failed if it exists
        if assistant_msg_id:
            try:
                assistant_msg = db.query(Message).filter(
                    Message.id == assistant_msg_id
                ).first()
                if assistant_msg:
                    assistant_msg.status = "failed"
                    assistant_msg.content = f"Error: An unexpected error occurred"
                    db.commit()
            except:
                pass
        
        yield json.dumps({
            "type": "error",
            "message": "An unexpected error occurred"
        }) + "\n"
        yield json.dumps({"type": "done"}) + "\n"


# Add to your database functions or services

from datetime import date
from sqlalchemy.orm import Session
from sqlalchemy import func

import razorpay
import os

# Initialize Razorpay
razorpay_client = razorpay.Client(
    auth=(os.getenv("RAZORPAY_KEY_ID"), os.getenv("RAZORPAY_KEY_SECRET"))
)

# Payment packages
PAYMENT_PACKAGES = {
    "starter": {
        "name": "Starter Pack",
        "credits": 200,
        "price": 4.99,
        "currency": "USD",
        "description": "200 messages"
    },
    "popular": {
        "name": "Popular Pack",
        "credits": 500,
        "price": 9.99,
        "currency": "USD",
        "description": "500 messages",
        "badge": "Most Popular"
    },
    "pro": {
        "name": "Pro Pack",
        "credits": 1200,
        "price": 19.99,
        "currency": "USD",
        "description": "1200 messages",
        "badge": "Best Deal"
    }
}

ANONYMOUS_DAILY_LIMIT = 1  # 2 free messages per day
LOGGED_IN_DAILY_LIMIT = 2

async def check_user_daily_limit(db: Session, user_id: str) -> tuple[bool, int]:
    """
    Check if user has exceeded daily limit.
    Returns: (is_allowed, remaining_messages)
    """
    today = date.today()
    
    # Get or create usage record for today
    usage = db.query(UserDailyUsage).filter(
        UserDailyUsage.user_id == user_id,
        UserDailyUsage.usage_date == today
    ).first()
    
    if not usage:
        # First message today
        return True, LOGGED_IN_DAILY_LIMIT
    
    remaining = LOGGED_IN_DAILY_LIMIT - usage.message_count
    is_allowed = usage.message_count < LOGGED_IN_DAILY_LIMIT
    
    return is_allowed, max(0, remaining)


async def increment_user_daily_usage(db: Session, user_id: str):
    """
    Increment message count for user.
    """
    today = date.today()
    
    usage = db.query(UserDailyUsage).filter(
        UserDailyUsage.user_id == user_id,
        UserDailyUsage.usage_date == today
    ).first()
    
    if usage:
        usage.message_count += 1
        usage.updated_at = datetime.utcnow()
    else:
        usage = UserDailyUsage(
            user_id=user_id,
            usage_date=today,
            message_count=1
        )
        db.add(usage)
    
    db.commit()
    
async def check_anonymous_limit(db: Session, ip_address: str) -> tuple[bool, int]:
    """
    Check if anonymous user has exceeded daily limit.
    Returns: (is_allowed, remaining_messages)
    """
    today = date.today()
    
    # Get or create usage record for today
    usage = db.query(AnonymousUsage).filter(
        AnonymousUsage.ip_address == ip_address,
        AnonymousUsage.usage_date == today
    ).first()
    
    if not usage:
        # First message today
        return True, ANONYMOUS_DAILY_LIMIT
    
    remaining = ANONYMOUS_DAILY_LIMIT - usage.message_count
    is_allowed = usage.message_count < ANONYMOUS_DAILY_LIMIT
    
    return is_allowed, max(0, remaining)


async def increment_anonymous_usage(db: Session, ip_address: str):
    """
    Increment message count for anonymous user.
    """
    today = date.today()
    
    usage = db.query(AnonymousUsage).filter(
        AnonymousUsage.ip_address == ip_address,
        AnonymousUsage.usage_date == today
    ).first()
    
    if usage:
        usage.message_count += 1
        usage.updated_at = datetime.utcnow()
    else:
        usage = AnonymousUsage(
            ip_address=ip_address,
            usage_date=today,
            message_count=1
        )
        db.add(usage)
    
    db.commit()

from fastapi import Request
def get_client_ip(request: Request) -> str:
    """
    Get the real client IP address, handling proxies and load balancers.
    """
    # Check for forwarded IP (if behind proxy/load balancer)
    forwarded = request.headers.get("X-Forwarded-For")
    if forwarded:
        # X-Forwarded-For can contain multiple IPs, take the first one
        return forwarded.split(",")[0].strip()
    
    # Check for real IP header
    real_ip = request.headers.get("X-Real-IP")
    if real_ip:
        return real_ip
    
    # Fallback to direct client IP
    if request.client:
        return request.client.host
    
    return "unknown"
        
@app.options("/auth/register")
async def register_options(request: Request):
    origin = request.headers.get("origin", "*")
    return Response(
        status_code=200,
        headers={
            "Access-Control-Allow-Origin": origin,
            "Access-Control-Allow-Credentials": "true",
            "Access-Control-Allow-Methods": "POST, OPTIONS",
            "Access-Control-Allow-Headers": "*",
        }
    )

@app.post("/auth/register")
async def register(
    request: Request,
    user_data: UserCreate,
    db: Session = Depends(get_db)
):
    existing = db.query(User).filter(
        (User.username == user_data.username) | (User.email == user_data.email)
    ).first()
    
    if existing:
        raise HTTPException(status_code=400, detail="Username or email already exists")
    
    hashed_password = pwd_context.hash(user_data.password)
    user = User(
        username=user_data.username,
        email=user_data.email,
        hashed_password=hashed_password
    )
    db.add(user)
    db.commit()
    db.refresh(user)
    
    token = create_access_token({"user_id": str(user.id), "username": user.username})
    
    # Get origin for CORS
    origin = request.headers.get("origin", "*")
    
    response_data = {
        "access_token": token,
        "token_type": "bearer",
        "user_id": str(user.id),
        "username": user.username
    }
    
    return Response(
        content=json.dumps(response_data),
        media_type="application/json",
        headers={
            "Access-Control-Allow-Origin": origin,
            "Access-Control-Allow-Credentials": "true",
        }
    )
     
@app.post("/auth/login")
async def login(user_data: UserLogin, db: Session = Depends(get_db)):
    user = db.query(User).filter(User.email == user_data.email).first()
    
    if not user or not pwd_context.verify(user_data.password, user.hashed_password):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    token = create_access_token({"user_id": str(user.id), "username": user.username})
    
    return {
        "access_token": token,
        "token_type": "bearer",
        "user_id": str(user.id),
        "username": user.username
    }

# Conversation Routes (keep for backward compatibility)
@app.post("/conversations", response_model=ConversationResponse)
async def create_conversation(
    conv_data: ConversationCreate,
    db: Session = Depends(get_db),
    current_user: Optional[dict] = Depends(get_current_user)
):
    conversation = Conversation(
        user_id=uuid.UUID(current_user["user_id"]) if current_user else None,
        title=conv_data.title,
        is_anonymous=current_user is None
    )
    db.add(conversation)
    db.commit()
    db.refresh(conversation)
    
    return ConversationResponse(
        id=str(conversation.id),
        title=conversation.title,
        created_at=conversation.created_at,
        updated_at=conversation.updated_at,
        message_count=0
    )

@app.get("/conversations", response_model=List[ConversationResponse])
async def list_conversations(
    db: Session = Depends(get_db),
    current_user: dict = Depends(verify_token)
):
    conversations = db.query(Conversation).filter(
        # Conversation.user_id == current_user["user_id"]
        Conversation.user_id == uuid.UUID(current_user["user_id"])
    ).order_by(Conversation.updated_at.desc()).all()
    
    return [
        ConversationResponse(
            id=str(conv.id),
            title=conv.title,
            created_at=conv.created_at,
            updated_at=conv.updated_at,
            message_count=conv.message_count or 0
        )
        for conv in conversations
    ]

@app.get("/conversations/{conversation_id}/messages")
async def get_messages(
    conversation_id: str,
    db: Session = Depends(get_db),
    current_user: Optional[dict] = Depends(get_current_user)
):
    conv = db.query(Conversation).filter(Conversation.id == uuid.UUID(conversation_id)).first()
    
    if not conv:
        raise HTTPException(status_code=404, detail="Conversation not found")
    
    if current_user and conv.user_id and str(conv.user_id) != current_user["user_id"]:
        raise HTTPException(status_code=403, detail="Not authorized")
    
    messages = db.query(Message).filter(
        Message.conversation_id == uuid.UUID(conversation_id)
    ).order_by(Message.created_at).all()
    
    result = []
    for m in messages:
        # Fetch reasoning steps from ReasoningStep table
        reasoning_steps = None
        if m.role == "assistant":
            # Query ReasoningStep table for this message
            steps = db.query(ReasoningStep).filter(
                ReasoningStep.message_id == m.id
            ).order_by(ReasoningStep.step_number).all()
            
            if steps:
                reasoning_steps = []
                for step in steps:
                    step_data = {
                        "step": step.content,
                        "content": step.content,
                        "step_number": step.step_number
                    }
                    
                    # Add optional fields if they exist
                    if step.query:
                        step_data["query"] = step.query
                    if step.category:
                        step_data["category"] = step.category
                    if step.sources:
                        try:
                            step_data["sources"] = json.loads(step.sources)
                        except:
                            step_data["sources"] = []
                    
                    reasoning_steps.append(step_data)
        
        result.append({
            "id": str(m.id),
            "role": m.role,
            "content": m.content,
            "has_file": m.has_file,
            "created_at": m.created_at.isoformat(),
            "sources": json.loads(m.sources) if m.sources else None,
            "reasoning_steps": reasoning_steps,  # From ReasoningStep table
            "assets": json.loads(m.assets) if m.assets else None,
            "app": m.app,
            "lab_mode": m.lab_mode,
             "status": m.status,
            "reactions": [{"type": r.reaction_type, "user_id": str(r.user_id)} for r in m.reactions]
        })
    
    return result
    
@app.delete("/conversations/{conversation_id}")
async def delete_conversation(
    conversation_id: str,
    db: Session = Depends(get_db),
    current_user: dict = Depends(verify_token)
):
    conv = db.query(Conversation).filter(
        Conversation.id == uuid.UUID(conversation_id),
        Conversation.user_id == uuid.UUID(current_user["user_id"])
    ).first()
    
    if not conv:
        raise HTTPException(status_code=404, detail="Conversation not found")
    
    db.delete(conv)
    db.commit()
    
    await conversation_manager.redis.delete(f"conv:{conversation_id}:history")
    
    return {"message": "Conversation deleted"}

@app.get("/health")
async def health_check():
    return {
        "status": "healthy",
        "timestamp": datetime.now(timezone.utc).isoformat(),
        # "spacy_loaded": nlp is not None,
        "features": {
            "ner_extraction": nlp is not None,
            "streaming": True,
            "anonymous_mode": True,
            "web_search": bool(SERPAPI_KEY)
        }
    }

@app.get("/messages/{message_id}/export/{format}")
async def export_message(
    message_id: str,
    format: str,
    db: Session = Depends(get_db),
    current_user: Optional[dict] = Depends(get_current_user)
):
    """Export message in PDF, DOCX, or MD format"""
    if format not in ['pdf', 'docx', 'md']:
        raise HTTPException(status_code=400, detail="Invalid format. Use pdf, docx, or md")
    
    message = db.query(Message).filter(Message.id == uuid.UUID(message_id)).first()
    
    if not message:
        raise HTTPException(status_code=404, detail="Message not found")
    
    # Parse sources
    sources = None
    if message.sources:
        try:
            sources = json.loads(message.sources) if isinstance(message.sources, str) else message.sources
        except:
            pass
    
    if format == 'pdf':
        buffer = await generate_message_pdf(message.content, sources)
        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f"attachment; filename=noir-ai-response-{message_id[:8]}.pdf",
                "Access-Control-Allow-Origin": "*", "Access-Control-Allow-Methods": "POST, OPTIONS", "Access-Control-Allow-Headers": "*"
            }
        )
    elif format == 'docx':
        buffer = await generate_message_docx(message.content, sources)
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename=noir-ai-response-{message_id[:8]}.docx",
                "Access-Control-Allow-Origin": "*", "Access-Control-Allow-Methods": "POST, OPTIONS", "Access-Control-Allow-Headers": "*"
            }
        )
    else:  # md
        md_content = await generate_message_markdown(message.content, sources)
        return Response(
            content=md_content,
            media_type="text/markdown",
            headers={
                "Content-Disposition": f"attachment; filename=noir-ai-response-{message_id[:8]}.md"
            }
        )
    
@app.get("/")
async def root():
    return FileResponse("coming_soon.html")


import asyncio
import json
import base64
import websockets
from websockets.server import serve
import os
import logging 

# Disable websockets debug logging
logging.getLogger('websockets').setLevel(logging.WARNING)
logging.getLogger('websockets.client').setLevel(logging.WARNING)
logging.getLogger('websockets.server').setLevel(logging.WARNING)
 
from fastapi import WebSocket, WebSocketDisconnect, Depends
from sqlalchemy.orm import Session
import asyncio
import json
import os
import logging
import websockets
from datetime import datetime
import uuid

logging.getLogger('websockets').setLevel(logging.WARNING)
logging.getLogger('websockets.client').setLevel(logging.WARNING)
logging.getLogger('websockets.server').setLevel(logging.WARNING)

class RealtimeVoiceHandler:
    def __init__(self, db: Session = None, user_id: str = None):
        self.openai_ws = None
        self.client_ws = None
        self.db = db
        self.user_id = user_id
        self.current_conversation_id = None
        self.user_transcript = ""
        self.assistant_transcript = ""
        self.assistant_content_complete = False
        
    async def connect_to_openai(self):
        """Establish WebSocket connection to OpenAI Realtime API"""
        url = "wss://api.openai.com/v1/realtime?model=gpt-4o-realtime-preview-2024-10-01"
        headers = {
            "Authorization": f"Bearer {os.getenv('OPENAI_API_KEY')}",
            "OpenAI-Beta": "realtime=v1"
        }
        
        print("[VOICE] Connecting to OpenAI Realtime API...")
        self.openai_ws = await websockets.connect(url, extra_headers=headers)
        print("[VOICE] ✅ Connected to OpenAI")
        await self.send_session_update()
        print("[VOICE] ✅ Session configured")
        
    async def send_session_update(self):
        """Configure the Realtime API session"""
        session_config = {
            "type": "session.update",
            "session": {
                "modalities": ["text", "audio"],
                "instructions": """You are a helpful AI assistant with web search capabilities. 
                When users ask questions that require current information or web data, use the search_web function.
                Always acknowledge when you're searching: say something like 'Let me search for that information' or 'I'll look that up for you'.
                
                IMPORTANT: When you receive search results from the search_web function, you MUST:
                1. Use the actual web content provided in the results
                2. Base your answer on the scraped information
                3. Cite the sources when appropriate
                4. Be specific and detailed using the data you received
                5. Do NOT make up information - only use what's in the search results
                
                When presenting search results, be conversational and natural while staying factual.""",
                "voice": "alloy",
                "input_audio_format": "pcm16",
                "output_audio_format": "pcm16",
                "input_audio_transcription": {
                    "model": "whisper-1"
                },
                "turn_detection": {
                    "type": "server_vad",
                    "threshold": 0.5,
                    "prefix_padding_ms": 300,
                    "silence_duration_ms": 500
                },
                "tools": [
                    {
                        "type": "function",
                        "name": "search_web",
                        "description": "Search the web for current information, news, data, or any query that requires real-time information.",
                        "parameters": {
                            "type": "object",
                            "properties": {
                                "query": {
                                    "type": "string",
                                    "description": "The search query to look up on the web"
                                }
                            },
                            "required": ["query"]
                        }
                    }
                ],
                "tool_choice": "auto"
            }
        }
        
        await self.openai_ws.send(json.dumps(session_config))
        
    async def handle_function_call(self, function_name, arguments, call_id):
        """Execute function calls from the AI"""
        from scraper_stable_optimized import search_with_web
        if function_name == "search_web":
            query = arguments.get("query", "")
            
            try:
                # Call your existing scraper
                full_text, urls, tables = await search_with_web(query)
                
                # Send visual data (tables/charts) ONLY to client for display
                if self.client_ws:
                    await self.client_ws.send_json({
                        "type": "search_results",
                        "query": query,
                        "urls": urls,
                        "tables": tables
                    })
                
                # Send ONLY text content to the LLM
                formatted_output = f"""Search Results for: {query}

Scraped Web Content:
{full_text}

Sources:
{chr(10).join(f"- {url}" for url in urls[:5])}

Instructions: Use the above scraped content to answer the user's question. Be specific and cite information from the content provided."""
                
                # Send function result back to OpenAI
                function_output = {
                    "type": "conversation.item.create",
                    "item": {
                        "type": "function_call_output",
                        "call_id": call_id,
                        "output": formatted_output
                    }
                }
                
                await self.openai_ws.send(json.dumps(function_output))
                await self.openai_ws.send(json.dumps({"type": "response.create"}))
                
            except Exception as e:
                error_output = {
                    "type": "conversation.item.create",
                    "item": {
                        "type": "function_call_output",
                        "call_id": call_id,
                        "output": json.dumps({"error": str(e)})
                    }
                }
                await self.openai_ws.send(json.dumps(error_output))
                
    async def handle_openai_messages(self):
        """Listen for messages from OpenAI and relay to client"""
        print("[VOICE] OpenAI message handler started")
        try:
            async for message in self.openai_ws:
                data = json.loads(message)
                event_type = data.get("type")
                
                # Only log important events, not every audio delta
                if event_type not in ["response.audio.delta"]:
                    print(f"[VOICE] OpenAI event: {event_type}")
                
                # Capture user transcript
                if event_type == "conversation.item.input_audio_transcription.completed":
                    self.user_transcript = data.get("transcript", "")
                    print(f"[VOICE] User transcript captured: {self.user_transcript[:50]}...")
                
                # Capture assistant transcript (streaming)
                if event_type == "response.audio_transcript.delta":
                    self.assistant_transcript += data.get("delta", "")
                
                # Assistant transcript complete
                if event_type == "response.audio_transcript.done":
                    print(f"[VOICE] Assistant transcript complete: {len(self.assistant_transcript)} chars")
                    self.assistant_content_complete = True
                
                # Handle function calls
                if event_type == "response.function_call_arguments.done":
                    function_name = data.get("name")
                    arguments = json.loads(data.get("arguments", "{}"))
                    call_id = data.get("call_id")
                    print(f"[VOICE] Function call: {function_name} with args: {arguments}")
                    await self.handle_function_call(function_name, arguments, call_id)
                    continue
                
                # Save to DB when response is complete
                if event_type == "response.done":
                    print("[VOICE] Response done - saving to DB")
                    await self.save_conversation_to_db()
                    # Reset for next interaction
                    self.user_transcript = ""
                    self.assistant_transcript = ""
                    self.assistant_content_complete = False
                
                # Relay audio and transcripts to client
                if event_type in [
                    "response.audio.delta",
                    "response.audio.done",
                    "response.audio_transcript.delta",
                    "response.audio_transcript.done",
                    "conversation.item.input_audio_transcription.completed",
                    "response.done",
                    "input_audio_buffer.speech_started",
                    "input_audio_buffer.speech_stopped"
                ]:
                    if self.client_ws:
                        await self.client_ws.send_text(message)
                        
        except websockets.exceptions.ConnectionClosed:
            print("[VOICE] OpenAI connection closed")
        except Exception as e:
            print(f"[VOICE] Error in OpenAI handler: {e}")
            import traceback
            traceback.print_exc()
    
    async def save_conversation_to_db(self):
        """Save voice conversation to database"""
        if not self.db:
            print("[VOICE] No database session available")
            return
        
        if not self.user_transcript or not self.assistant_transcript:
            print(f"[VOICE] Missing transcripts - user: {bool(self.user_transcript)}, assistant: {bool(self.assistant_transcript)}")
            return
        
        try:
            # Create or get conversation
            if not self.current_conversation_id:
                # Create new conversation with a title from first user message
                title = self.user_transcript[:50] + ("..." if len(self.user_transcript) > 50 else "")
                
                new_conversation = Conversation(
                    user_id=uuid.UUID(self.user_id) if self.user_id else None,
                    title=title,
                    created_at=datetime.utcnow(),
                    updated_at=datetime.utcnow()
                )
                self.db.add(new_conversation)
                self.db.flush()  # Get the ID without committing
                self.current_conversation_id = new_conversation.id
                print(f"[VOICE] Created new conversation: {self.current_conversation_id}")
                
                # Send conversation ID to client
                if self.client_ws:
                    await self.client_ws.send_json({
                        "type": "conversation_created",
                        "conversation_id": str(self.current_conversation_id)
                    })
            
            # Save user message
            user_msg = Message(
                conversation_id=uuid.UUID(str(self.current_conversation_id)),
                role="user",
                content=self.user_transcript,
                has_file=False,
                created_at=datetime.utcnow()
            )
            self.db.add(user_msg)
            
            # Save assistant message
            assistant_msg = Message(
                conversation_id=uuid.UUID(str(self.current_conversation_id)),
                role="assistant",
                content=self.assistant_transcript,
                created_at=datetime.utcnow()
            )
            self.db.add(assistant_msg)
            
            # Update conversation timestamp
            conversation = self.db.query(Conversation).filter(
                Conversation.id == uuid.UUID(str(self.current_conversation_id))
            ).first()
            if conversation:
                conversation.updated_at = datetime.utcnow()
            
            self.db.commit()
            print(f"[VOICE] ✅ Saved conversation to DB (user: {len(self.user_transcript)} chars, assistant: {len(self.assistant_transcript)} chars)")
            
        except Exception as e:
            print(f"[VOICE] ❌ Failed to save to DB: {e}")
            import traceback
            traceback.print_exc()
            self.db.rollback()
    
    async def handle_client_messages(self):
        try:
            while True:
                message = await self.client_ws.receive_text()
                msg_data = json.loads(message)
                msg_type = msg_data.get("type", "unknown")
                
                # When user starts speaking, cancel any ongoing response
                if msg_type == "input_audio_buffer.speech_started":
                    cancel_msg = {"type": "response.cancel"}
                    await self.openai_ws.send(json.dumps(cancel_msg))
                
                if msg_type == "input_audio_buffer.append":
                    pass  # Don't log every audio chunk
                else:
                    print(f"[VOICE] Client message: {msg_type}")
                
                # Forward to OpenAI
                await self.openai_ws.send(message)
                
        except WebSocketDisconnect:
            print("[VOICE] Client disconnected")
        except Exception as e:
            print(f"[VOICE] Error in client handler: {e}")
            import traceback
            traceback.print_exc()
        
    async def connect_to_openai(self):
        url = "wss://api.openai.com/v1/realtime?model=gpt-4o-realtime-preview-2024-10-01"
        headers = {
            "Authorization": f"",
            "OpenAI-Beta": "realtime=v1"
        }
        
        self.openai_ws = await websockets.connect(url, extra_headers=headers)
        await self.send_session_update()
        
    async def send_session_update(self):
        session_config = {
            "type": "session.update",
            "session": {
                "modalities": ["text", "audio"],
                "instructions": """You are a helpful AI assistant with web search capabilities. 
                When users ask questions that require current information or web data, use the search_web function.
                Always acknowledge when you're searching: say something like 'Let me search for that information' or 'I'll look that up for you'.
                
                IMPORTANT: When you receive search results from the search_web function, you MUST:
                1. Use the actual web content provided in the results
                2. Base your answer on the scraped information
                3. Cite the sources when appropriate
                4. Be specific and detailed using the data you received
                5. Do NOT make up information - only use what's in the search results
                
                When presenting search results, be conversational and natural while staying factual.""",
                "voice": "alloy",
                "input_audio_format": "pcm16",
                "output_audio_format": "pcm16",
                "input_audio_transcription": {
                    "model": "whisper-1"
                },
                "turn_detection": {
                    "type": "server_vad",
                    "threshold": 0.5,
                    "prefix_padding_ms": 300,
                    "silence_duration_ms": 500
                },
                "tools": [
                    {
                        "type": "function",
                        "name": "search_web",
                        "description": "Search the web for current information, news, data, or any query that requires real-time information.",
                        "parameters": {
                            "type": "object",
                            "properties": {
                                "query": {
                                    "type": "string",
                                    "description": "The search query to look up on the web"
                                }
                            },
                            "required": ["query"]
                        }
                    }
                ],
                "tool_choice": "auto"
            }
        }
        
        await self.openai_ws.send(json.dumps(session_config))
        
    async def handle_function_call(self, function_name, arguments, call_id):
        if function_name == "search_web":
            query = arguments.get("query", "")
            from scraper_stable_optimized import search_with_web
            try:
                full_text, urls, tables = await search_with_web(query)
                
                if self.client_ws:
                    await self.client_ws.send_json({
                        "type": "search_results",
                        "query": query,
                        "urls": urls,
                        "tables": tables
                    })
                
                # Send ONLY text content to the LLM
                formatted_output = f"""Search Results for: {query}

Scraped Web Content:
{full_text}

Sources:
{chr(10).join(f"- {url}" for url in urls[:5])}

Instructions: Use the above scraped content to answer the user's question. Be specific and cite information from the content provided."""
                
                # Send function result back to OpenAI
                function_output = {
                    "type": "conversation.item.create",
                    "item": {
                        "type": "function_call_output",
                        "call_id": call_id,
                        "output": formatted_output
                    }
                }
                
                await self.openai_ws.send(json.dumps(function_output))
                await self.openai_ws.send(json.dumps({"type": "response.create"}))
                
            except Exception as e:
                error_output = {
                    "type": "conversation.item.create",
                    "item": {
                        "type": "function_call_output",
                        "call_id": call_id,
                        "output": json.dumps({"error": str(e)})
                    }
                }
                await self.openai_ws.send(json.dumps(error_output))
                
    async def handle_openai_messages(self):
        """Listen for messages from OpenAI and relay to client"""
        print("[VOICE] OpenAI message handler started")
        try:
            async for message in self.openai_ws:
                data = json.loads(message)
                event_type = data.get("type")
                
                # Only log important events, not every audio delta
                if event_type not in ["response.audio.delta"]:
                    print(f"[VOICE] OpenAI event: {event_type}")
                
                # Capture user transcript
                if event_type == "conversation.item.input_audio_transcription.completed":
                    self.user_transcript = data.get("transcript", "")
                    print(f"[VOICE] User transcript captured: {self.user_transcript[:50]}...")
                
                # Capture assistant transcript (streaming)
                if event_type == "response.audio_transcript.delta":
                    self.assistant_transcript += data.get("delta", "")
                
                # Assistant transcript complete
                if event_type == "response.audio_transcript.done":
                    print(f"[VOICE] Assistant transcript complete: {len(self.assistant_transcript)} chars")
                    self.assistant_content_complete = True
                
                # Handle function calls
                if event_type == "response.function_call_arguments.done":
                    function_name = data.get("name")
                    arguments = json.loads(data.get("arguments", "{}"))
                    call_id = data.get("call_id")
                    print(f"[VOICE] Function call: {function_name} with args: {arguments}")
                    await self.handle_function_call(function_name, arguments, call_id)
                    continue
                
                # Save to DB when response is complete
                if event_type == "response.done":
                    print("[VOICE] Response done - saving to DB")
                    await self.save_conversation_to_db()
                    # Reset for next interaction
                    self.user_transcript = ""
                    self.assistant_transcript = ""
                    self.assistant_content_complete = False
                
                # Relay audio and transcripts to client
                if event_type in [
                    "response.audio.delta",
                    "response.audio.done",
                    "response.audio_transcript.delta",
                    "response.audio_transcript.done",
                    "conversation.item.input_audio_transcription.completed",
                    "response.done",
                    "input_audio_buffer.speech_started",
                    "input_audio_buffer.speech_stopped"
                ]:
                    if self.client_ws:
                        await self.client_ws.send_text(message)
                        
        except websockets.exceptions.ConnectionClosed:
            print("[VOICE] OpenAI connection closed")
        except Exception as e:
            print(f"[VOICE] Error in OpenAI handler: {e}")
            import traceback
            traceback.print_exc()
    
    async def save_conversation_to_db(self):
        """Save voice conversation to database"""
        if not self.db:
            print("[VOICE] No database session available")
            return
        
        if not self.user_transcript or not self.assistant_transcript:
            print("[VOICE] Missing transcripts - not saving")
            return
        
        try:
            # Create or get conversation
            if not self.current_conversation_id:
                # Create new conversation with a title from first user message
                title = self.user_transcript[:50] + ("..." if len(self.user_transcript) > 50 else "")
                
                new_conversation = Conversation(
                    user_id=uuid.UUID(self.user_id) if self.user_id else None,
                    title=title,
                    created_at=datetime.utcnow(),
                    updated_at=datetime.utcnow()
                )
                self.db.add(new_conversation)
                self.db.flush()  # Get the ID without committing
                self.current_conversation_id = new_conversation.id
                print(f"[VOICE] Created new conversation: {self.current_conversation_id}")
                
                # Send conversation ID to client
                if self.client_ws:
                    await self.client_ws.send_json({
                        "type": "conversation_created",
                        "conversation_id": str(self.current_conversation_id)
                    })
            
            # Save user message
            user_msg = Message(
                conversation_id=uuid.UUID(str(self.current_conversation_id)),
                role="user",
                content=self.user_transcript,
                has_file=False,
                created_at=datetime.utcnow()
            )
            self.db.add(user_msg)
            
            # Save assistant message
            assistant_msg = Message(
                conversation_id=uuid.UUID(str(self.current_conversation_id)),
                role="assistant",
                content=self.assistant_transcript,
                created_at=datetime.utcnow()
            )
            self.db.add(assistant_msg)
            
            # Update conversation timestamp
            conversation = self.db.query(Conversation).filter(
                Conversation.id == uuid.UUID(str(self.current_conversation_id))
            ).first()
            if conversation:
                conversation.updated_at = datetime.utcnow()
            
            self.db.commit()
            print(f"[VOICE] ✅ Saved conversation to DB")
            
        except Exception as e:
            print(f"[VOICE] ❌ Failed to save to DB: {e}")
            import traceback
            traceback.print_exc()
            self.db.rollback()
    
    async def handle_client_messages(self):
        """Listen for messages from browser client and relay to OpenAI"""
        print("[VOICE] Client message handler started")
        try:
            while True:
                message = await self.client_ws.receive_text()
                msg_data = json.loads(message)
                msg_type = msg_data.get("type", "unknown")
                
                # When user starts speaking, cancel any ongoing response
                if msg_type == "input_audio_buffer.speech_started":
                    print("[VOICE] User started speaking - canceling response")
                    cancel_msg = {"type": "response.cancel"}
                    await self.openai_ws.send(json.dumps(cancel_msg))
                
                if msg_type == "input_audio_buffer.append":
                    pass  # Don't log every audio chunk
                else:
                    print(f"[VOICE] Client message: {msg_type}")
                
                # Forward to OpenAI
                await self.openai_ws.send(message)
                
        except WebSocketDisconnect:
            print("[VOICE] Client disconnected")
        except Exception as e:
            print(f"[VOICE] Error in client handler: {e}")
            import traceback
            traceback.print_exc()


@app.websocket("/ws/voice")
async def voice_websocket(websocket: WebSocket, db: Session = Depends(get_db)):
    """
    FastAPI WebSocket endpoint for voice interface with Realtime API
    
    Query Parameters:
        - token (optional): JWT authentication token
    
    Features:
        - Real-time voice conversation with OpenAI
        - Automatic transcript capture
        - Database persistence
        - Web search integration
    """
    print("[VOICE] New WebSocket connection attempt")
    await websocket.accept()
    print("[VOICE] WebSocket connection accepted")
    
    # Get user from token if authenticated
    user_id = None
    try:
        token = websocket.query_params.get("token")
        if token:
            # Verify token and get user_id
            # UPDATE THIS to match your auth implementation
            from jose import jwt, JWTError
            # from your_auth_file import SECRET_KEY, ALGORITHM  # Your JWT settings
            
            try:
                payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
                user_id = payload.get("sub")
                print(f"[VOICE] Authenticated user: {user_id}")
            except JWTError:
                print("[VOICE] Invalid token, continuing as guest")
    except Exception as e:
        print(f"[VOICE] No authentication: {e}")
    
    # Create handler with database and user context
    handler = RealtimeVoiceHandler(db=db, user_id=user_id)
    handler.client_ws = websocket
    
    try:
        # Connect to OpenAI
        await handler.connect_to_openai()
        
        print("[VOICE] Starting message relay loops...")
        
        # Start bidirectional message relay
        await asyncio.gather(
            handler.handle_openai_messages(),
            handler.handle_client_messages()
        )
        
    except WebSocketDisconnect:
        print("[VOICE] WebSocket disconnected")
    except Exception as e:
        print(f"[VOICE] Error in voice handler: {e}")
        import traceback
        traceback.print_exc()
    finally:
        if handler.openai_ws:
            await handler.openai_ws.close()
        print("[VOICE] Connection closed")

@app.get("/voice/health")
async def voice_health():
    """Check if voice service is available"""
    return {
        "status": "ok",
        "service": "voice",
        "openai_api_key_set": bool(os.getenv("OPENAI_API_KEY"))
    }
     
@app.get("/debug/config")
async def debug_config():
    """Remove this endpoint in production!"""
    return {
        "client_id": GOOGLE_CLIENT_ID[:20] + "..." if GOOGLE_CLIENT_ID else "NOT SET",
        "client_secret": "SET" if GOOGLE_CLIENT_SECRET else "NOT SET",
        "redirect_uri": REDIRECT_URI
    }

import os
from dotenv import load_dotenv
from fastapi import FastAPI, Request, Depends, HTTPException
from fastapi.responses import RedirectResponse
from sqlalchemy.orm import Session
import requests
from jose import jwt
from datetime import datetime, timedelta
import uuid

@app.get("/auth/login")
async def google_login():
    """Redirect to Google OAuth"""
    google_auth_url = (
        f"https://accounts.google.com/o/oauth2/v2/auth"
        f"?client_id={GOOGLE_CLIENT_ID}"
        f"&redirect_uri={REDIRECT_URI}"  # Backend callback URL
        f"&response_type=code"
        f"&scope=openid email profile"
        f"&access_type=offline"
    )
    
    return RedirectResponse(url=google_auth_url)
  
@app.get("/auth/callback")
async def google_callback(code: str, db: Session = Depends(get_db)):
    """Handle Google OAuth callback"""
    try:
        if not code:
            return RedirectResponse(url=f"{FRONTEND_URL}?error=no_code", status_code=302)
        
        token_url = "https://oauth2.googleapis.com/token"
        token_data = {
            "code": code,
            "client_id": GOOGLE_CLIENT_ID,
            "client_secret": GOOGLE_CLIENT_SECRET,
            "redirect_uri": REDIRECT_URI,
            "grant_type": "authorization_code"
        }
        
        token_response = requests.post(token_url, data=token_data)
        token_json = token_response.json()
        
        if "error" in token_json:
            print(f"❌ Token error: {token_json}")
            return RedirectResponse(url=f"{FRONTEND_URL}?error=token_failed", status_code=302)
        
        google_access_token = token_json.get("access_token")
        
        # Get user info
        user_info_response = requests.get(
            "https://www.googleapis.com/oauth2/v2/userinfo",
            headers={"Authorization": f"Bearer {google_access_token}"}
        )
        user_info = user_info_response.json()
        
        email = user_info.get("email")
        google_id = user_info.get("id")
        name = user_info.get("name", email)
        picture = user_info.get("picture")
        
        # Find or create user
        user = db.query(User).filter(User.email == email).first()
        
        if not user:
            user = User(
                id=uuid.uuid4(),
                email=email,
                username=name or email.split('@')[0],
                hashed_password=None,
                oauth_provider="google",
                oauth_id=google_id,
                google_name=name,
                google_picture=picture,
                created_at=datetime.utcnow(),
                last_login=datetime.utcnow()
            )
            db.add(user)
            db.commit()
            db.refresh(user)
            print(f"✅ Created user: {email}")
        else:
            user.oauth_provider = "google"
            user.oauth_id = google_id
            user.google_name = name
            user.google_picture = picture
            user.last_login = datetime.utcnow()
            db.commit()
            print(f"✅ Updated user: {email}")
        
        user_data = {
            "user_id": str(user.id),  # Convert UUID to string
            "email": user.email,
            "username": user.username,
            "name": user.google_name or user.username,
            "pic": user.google_picture
        }
        # Create JWT
        jwt_token = create_access_token(user_data)
        
        # ✅ SOLUTION: Pass token in URL instead of cookie
        redirect_url = (
            f"{FRONTEND_URL}?oauth_success=true"
            f"&token={jwt_token}"
            f"&user_id={str(user.id)}"
            f"&username={user.google_name or user.username}"
            f"&email={user.email}"
        )
        
        print(f"✅ OAuth successful for {email}")
        print(f"🔄 Redirecting to: {redirect_url[:100]}...")
        
        return RedirectResponse(url=redirect_url, status_code=302)
        
    except Exception as e:
        print(f"❌ OAuth error: {e}")
        import traceback
        traceback.print_exc()
        return RedirectResponse(url=f"{FRONTEND_URL}?error=auth_failed", status_code=302)
    
@app.get("/me")
async def get_me(request: Request):
    """Get current user info from session cookie"""
    token = request.cookies.get("session")
    
    if not token:
        raise HTTPException(status_code=401, detail="Not logged in")
    
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=["HS256"])
        
        return {
            "user_id": payload.get("user_id"),
            # "email": payload["email"],
            "username": payload.get("username"),
            # "name": payload.get("name"),
            # "picture": payload.get("pic")
        }
    except Exception as e:
        print(f"❌ Token decode error: {e}")
        raise HTTPException(status_code=401, detail="Invalid token")

# class NewsArticle(BaseModel):
#     title: str
#     description: Optional[str] = None
#     link: str
#     pubDate: str
#     thumbnail: Optional[str] = None
#     source: str
#     author: Optional[str] = None

# class NewsResponse(BaseModel):
#     status: str
#     category: str
#     country: str
#     articles: List[NewsArticle]
#     cached_at: str
#     next_update: str
#     total_results: int

# Cache storage
news_cache: Dict[str, Dict] = {}
CACHE_DURATION = 300  # 5 minutes

background_task_running = False

# NewsData.io country codes (ISO 3166-1 alpha-2)
SUPPORTED_COUNTRIES = {
    "US": "us",
    "GB": "gb",
    "IN": "in",
    "AU": "au",
    "CA": "ca",
    "DE": "de",
    "FR": "fr",
    "IT": "it",
    "ES": "es",
    "NL": "nl",
    "NO": "no",
    "SE": "se",
    "RU": "ru",
    "CN": "cn",
    "JP": "jp",
    "KR": "kr",
    "BR": "br",
    "MX": "mx",
    "AR": "ar"
}

# Language codes for countries
COUNTRY_LANGUAGES = {
    "US": "en",
    "GB": "en",
    "IN": "en",
    "AU": "en",
    "CA": "en",
    "DE": "de",
    "FR": "fr",
    "IT": "it",
    "ES": "es",
    "NL": "nl",
    "NO": "no",
    "SE": "sv",
    "RU": "ru",
    "CN": "zh",
    "JP": "ja",
    "KR": "ko",
    "BR": "pt",
    "MX": "es",
    "AR": "es"
}

# NewsData.io category mapping
CATEGORY_MAP = {
    "general": "top",
    "sports": "sports",
    "technology": "technology",
    "business": "business",
    "entertainment": "entertainment",
    "politics": "politics"
}

async def detect_user_country(ip_address: str = None) -> str:
    """Detect user's country from IP address"""
    try:
        if ip_address:
            url = f"https://ipapi.co/{ip_address}/json/"
        else:
            url = "https://ipapi.co/json/"
        
        async with httpx.AsyncClient(timeout=10.0) as client:
            response = await client.get(url)
            if response.status_code == 200:
                data = response.json()
                country_code = data.get("country_code", "US")
                country_name = data.get("country_name", "Unknown")
                
                return country_code
            else:
                logger.warning(f"⚠️ IP detection returned status {response.status_code}")
    except Exception as e:
        logger.warning(f"⚠️ Could not detect country from IP: {e}")
    
    return "US"

async def fetch_news_from_newsdata(
    category: str = "general",
    country: str = "US",
    query: Optional[str] = None
) -> tuple[List[Dict], int]:
 
    try:
        if NEWSAPI_KEY == "YOUR_API_KEY_HERE":
            logger.error("❌ NewsData.io API key not set! Get one from https://newsdata.io")
            return [], 0
        
        all_articles = []
        base_url = "https://newsdata.io/api/1/news"
        
        newsdata_category = CATEGORY_MAP.get(category, "top")
        
        country_code = SUPPORTED_COUNTRIES.get(country.upper(), "us")
        
        language = COUNTRY_LANGUAGES.get(country.upper(), "en")
        
        # Build parameters
        params = {
            "apikey": NEWSAPI_KEY,
            "country": country_code,
            "language": language,
            "category": newsdata_category
        }
        
        # Add search query if provided
        if query:
            params["q"] = query
        
        async with httpx.AsyncClient(timeout=30.0) as client:
            response = await client.get(base_url, params=params)
            
            if response.status_code == 200:
                data = response.json()
                
                if data.get("status") == "success":
                    results = data.get("results", [])
                    total_results = data.get("totalResults", len(results))
                    
                    for item in results:
                        # Skip articles without required fields
                        if not item.get("title") or not item.get("link"):
                            continue
                        
                        # Get author(s) - NewsData.io returns list of creators
                        creators = item.get("creator")
                        author = None
                        if creators and isinstance(creators, list) and len(creators) > 0:
                            author = ", ".join(creators[:2])  # Take first 2 authors
                        
                        # Get source name
                        source_name = item.get("source_id", "Unknown").replace("_", " ").title()
                        
                        all_articles.append({
                            "title": item.get("title", "Untitled"),
                            "description": item.get("description") or item.get("content", "")[:200],
                            "link": item.get("link", ""),
                            "pubDate": item.get("pubDate", datetime.now().isoformat()),
                            "thumbnail": item.get("image_url"),
                            "source": source_name,
                            "author": author
                        })
                    
                    # logger.info(f"✅ NewsData.io: {len(all_articles)} articles fetched")
                    return all_articles, total_results
                else:
                    error_msg = data.get("message", "Unknown error")
                    logger.error(f"❌ NewsData.io API error: {error_msg}")
                    return [], 0
            
            elif response.status_code == 401:
                logger.error("❌ Invalid NewsData.io API key!")
                return [], 0
            elif response.status_code == 429:
                logger.error("❌ NewsData.io rate limit exceeded!")
                return [], 0
            else:
                logger.error(f"❌ NewsData.io HTTP error: {response.status_code}")
                logger.error(f"Response: {response.text}")
                return [], 0
                
    except httpx.HTTPStatusError as e:
        logger.error(f"❌ HTTP error: {e.response.status_code}")
        return [], 0
    except Exception as e:
        logger.error(f"❌ Error fetching news: {str(e)}")
        return [], 0


def get_cache_key(category: str, country: str, query: Optional[str] = None) -> str:
    """Generate cache key"""
    if query:
        return f"SEARCH_{query}_{country}".upper()
    return f"{category}_{country}".upper()

def is_cache_valid(cache_key: str) -> bool:
    """Check if cache is still valid"""
    if cache_key not in news_cache:
        return False
    
    cached_time = datetime.fromisoformat(news_cache[cache_key]["cached_at"])
    return datetime.now() - cached_time < timedelta(seconds=CACHE_DURATION)


async def update_cache(category: str, country: str, query: Optional[str] = None):
    """Update cache for a specific category and country"""
    cache_key = get_cache_key(category, country, query)
    
    articles, total_results = await fetch_news_from_newsdata(category, country, query)
    
    now = datetime.now()
    next_update = now + timedelta(seconds=CACHE_DURATION)
    
    news_cache[cache_key] = {
        "category": category,
        "country": country,
        "query": query,
        "articles": articles,
        "total_results": total_results,
        "cached_at": now.isoformat(),
        "next_update": next_update.isoformat()
    }
    
    if articles:
        logger.info(f"📦 Cache updated for {cache_key} - {len(articles)} articles")
    else:
        logger.warning(f"⚠️ No articles for {cache_key} - cached empty result")


async def background_cache_updater():
    """Background task to update cache every 5 minutes"""
    categories = ["general", "sports", "technology", "business", "entertainment", "politics"]
    countries = ["US", "AU", "GB", "IN", "CA"]  # Cache multiple countries
    
    while True:
        try:
            tasks = []
            for category in categories:
                for country in countries:
                    tasks.append(update_cache(category, country))
            
            await asyncio.gather(*tasks, return_exceptions=True)
            
        except Exception as e:
            logger.error(f"❌ Error in background cache updater: {str(e)}")
        
        await asyncio.sleep(CACHE_DURATION)

@app.on_event("startup")
async def startup_event():
    """Start background cache updater on app startup"""
    global background_task_running
    
    if not background_task_running:
        background_task_running = True
        asyncio.create_task(background_cache_updater())
 
@app.get("/news", response_model=NewsResponse)
async def get_news(
    category: str = "general",
    country: Optional[str] = None,
    q: Optional[str] = None,
    auto_detect: bool = True
):
    try:
        if not country and auto_detect:
            country = await detect_user_country()
        elif not country:
            country = "US"
        
        country = country.upper()
        cache_key = get_cache_key(category, country, q)
        
        if is_cache_valid(cache_key):
            # logger.info(f"📦 Serving from cache: {cache_key}")
            cached_data = news_cache[cache_key]
            
            return NewsResponse(
                status="ok",
                category=category,
                country=country,
                articles=[NewsArticle(**article) for article in cached_data["articles"]],
                cached_at=cached_data["cached_at"],
                next_update=cached_data["next_update"],
                total_results=cached_data.get("total_results", len(cached_data["articles"]))
            )
        
        # logger.info(f"🔍 Cache miss for {cache_key}, fetching fresh data...")
        await update_cache(category, country, q)
        
        if cache_key in news_cache:
            cached_data = news_cache[cache_key]
            return NewsResponse(
                status="ok",
                category=category,
                country=country,
                articles=[NewsArticle(**article) for article in cached_data["articles"]],
                cached_at=cached_data["cached_at"],
                next_update=cached_data["next_update"],
                total_results=cached_data.get("total_results", len(cached_data["articles"]))
            )
        else:
            return NewsResponse(
                status="ok",
                category=category,
                country=country,
                articles=[],
                cached_at=datetime.now().isoformat(),
                next_update=(datetime.now() + timedelta(seconds=CACHE_DURATION)).isoformat(),
                total_results=0
            )
            
    except Exception as e:
        logger.error(f"❌ Error in get_news: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/detect-country")
async def detect_country():
    """Detect user's country from IP address"""
    try:
        country = await detect_user_country()
        
        country_names = {
            "US": "United States",
            "GB": "United Kingdom", 
            "IN": "India",
            "AU": "Australia",
            "CA": "Canada",
            "DE": "Germany",
            "FR": "France",
            "IT": "Italy"
        }
        
        country_name = country_names.get(country, country)
        language = COUNTRY_LANGUAGES.get(country, "en")
        
        return {
            "country_code": country,
            "country_name": country_name,
            "language": language,
            "status": "detected"
        }
    except Exception as e:
        logger.error(f"❌ Detection error: {e}")
        return {
            "country_code": "US",
            "country_name": "United States",
            "language": "en",
            "status": "fallback"
        }

import secrets
from pydantic import EmailStr
from email_service import EmailService
        
class MagicLinkRequest(BaseModel):
    email: EmailStr
        
@app.post("/magic-link")
async def request_magic_link(
    request: MagicLinkRequest,
    req: Request,
    db: Session = Depends(get_db)
):
    """Request a magic link login email"""
    email = request.email.lower().strip()
    
    # Check if user exists, create if not
    user = db.query(User).filter(User.email == email).first()
    
    if not user:
        # Auto-create user account
        username = email.split('@')[0]
        base_username = username
        counter = 1
        while db.query(User).filter(User.username == username).first():
            username = f"{base_username}{counter}"
            counter += 1
        
        user = User(
            id=uuid.uuid4(),
            username=username,
            email=email,
            hashed_password=None,
            oauth_provider="magic_link",
            is_active=True
        )
        db.add(user)
        db.commit()
        db.refresh(user)
    
    # Generate secure token
    token = secrets.token_urlsafe(32)
    expires_at = datetime.now(timezone.utc) + timedelta(minutes=15)
    client_ip = req.client.host if req.client else None
    
    # Create magic link
    magic_link = MagicLink(
        email=email,
        token=token,
        expires_at=expires_at,
        ip_address=client_ip
    )
    db.add(magic_link)
    db.commit()
    
    # Send email
    email_sent = await EmailService.send_magic_link(email, token, user.username)
    
    if not email_sent:
        raise HTTPException(status_code=500, detail="Failed to send email")
    
    return {
        "message": "Magic link sent! Check your email.",
        "email": email,
        "expires_in_minutes": 15
    }

from fastapi.responses import RedirectResponse

@app.get("/verify")
async def verify_magic_link(
    token: str,
    db: Session = Depends(get_db)
):
    """Verify magic link and redirect to frontend"""
    last_record = db.query(MagicLink).order_by(MagicLink.created_at.desc()).first()
    
    if last_record:
        print("Last record ->" + str(last_record.token))
        print("Last record time ->" + str(last_record.used))
    else:
        print("no last record")
        
    print("\n\n TOKEN IN VERIFY \n\n" + token)
    # Verify token (same logic as before)
    magic_link = db.query(MagicLink).filter(
        MagicLink.token == token,
        MagicLink.used.is_(False)
    ).first()
         
    if magic_link:
        print("magic-link " + str(magic_link) )
        
        if datetime.now(timezone.utc) > magic_link.expires_at:
            print("expired")
    else:
        print("magic link is null")
            
    if not magic_link or datetime.now(timezone.utc) > magic_link.expires_at:
        # Redirect to frontend with error
        frontend_url = os.getenv("FRONTEND_URL")
        print("UN-SUCCESSFUL")
        return RedirectResponse(url=f"{frontend_url}/?error=invalid_link")
     
    # Mark as used
    magic_link.used = True
    magic_link.used_at = datetime.now(timezone.utc)
    
    # Get user
    user = db.query(User).filter(User.email == magic_link.email).first()
    user.last_login = datetime.now(timezone.utc)
    db.commit()
    
    # Generate JWT
    # from auth import create_access_token
    # access_token = create_access_token(data={"sub": str(user.id)})
    access_token =  create_access_token({"user_id": str(user.id), "username": user.username})
    # Redirect to frontend with credentials
    frontend_url = os.getenv("FRONTEND_URL")
    redirect_url = (
        f"{frontend_url}/"
        f"?oauth_success=true"
        f"&token={access_token}"
        f"&user_id={user.id}"
        f"&username={user.username}"
        f"&email={user.email}"
    )
    
    print("SUCCESSFUL")
    
    return RedirectResponse(url=redirect_url)

from use_cases_schema import *
@app.get("/use_cases")
async def get_all_use_cases(
    category: Optional[str] = None,
    featured: Optional[bool] = None,
    db: Session = Depends(get_db)
):
    try:
        query = db.query(PublicUseCase)
        
        # Apply filters
        if category:
            query = query.filter(PublicUseCase.category == category)
        
        if featured is not None:
            query = query.filter(PublicUseCase.featured == featured)
        
        # Order by featured first, then by views, then by created date
        query = query.order_by(
            PublicUseCase.featured.desc(),
            PublicUseCase.view_count.desc(),
            PublicUseCase.created_at.desc()
        )
        
        use_cases = query.all()
        
        return {
            "status": "success",
            "count": len(use_cases),
            "use_cases": [uc.to_dict() for uc in use_cases]
        }
        
    except Exception as e:
        print(f"❌ Error fetching use cases: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/use_cases/{use_case_id}/messages")
async def get_use_case_messages(
    use_case_id: str,
    db: Session = Depends(get_db)
):
    try:
        # Get the use case
        use_case = db.query(PublicUseCase).filter(
            PublicUseCase.id == use_case_id
        ).first()
        
        if not use_case:
            raise HTTPException(status_code=404, detail="Use case not found")
        
        # Get messages (already ordered by 'order' field via relationship)
        messages = use_case.messages
        
        return {
            "status": "success",
            "use_case": use_case.to_dict(),
            "messages": [msg.to_dict() for msg in messages]
        }
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Error fetching use case messages: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/use_cases/{use_case_id}/increment-views")
async def increment_use_case_views(
    use_case_id: str,
    db: Session = Depends(get_db)
):
    try:
        use_case = db.query(PublicUseCase).filter(
            PublicUseCase.id == use_case_id
        ).first()
        
        if not use_case:
            raise HTTPException(status_code=404, detail="Use case not found")
        
        # Increment view count
        use_case.view_count += 1
        db.commit()
        
        return {
            "status": "success",
            "use_case_id": str(use_case_id),
            "new_view_count": use_case.view_count
        }
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Error incrementing views: {e}")
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))


# ============================================
# GET CATEGORIES WITH COUNTS
# ============================================

@app.get("/use_cases/categories/list")
async def get_categories_with_counts(db: Session = Depends(get_db)):
    try:
        from sqlalchemy import func
        
        # Query to count use cases per category
        category_counts = db.query(
            PublicUseCase.category,
            func.count(PublicUseCase.id).label('count')
        ).group_by(PublicUseCase.category).all()
        
        # Format response
        categories = []
        for category, count in category_counts:
            category_info = USE_CASE_CATEGORIES.get(category, {
                'icon': '📁',
                'name': category.title()
            })
            categories.append({
                "key": category,
                "name": category_info['name'],
                "icon": category_info['icon'],
                "count": count
            })
        
        # Add "All" category
        total_count = sum(c['count'] for c in categories)
        categories.insert(0, {
            "key": "all",
            "name": "All Use Cases",
            "icon": "🌐",
            "count": total_count
        })
        
        return {
            "status": "success",
            "categories": categories
        }
        
    except Exception as e:
        print(f"❌ Error fetching categories: {e}")
        raise HTTPException(status_code=500, detail=str(e))


import razorpay
from fastapi import HTTPException
import os

# Initialize Razorpay client
razorpay_client = razorpay.Client(
    auth=(os.getenv("RAZORPAY_KEY_ID"), os.getenv("RAZORPAY_KEY_SECRET"))
)


# GET available packages
@app.get("/payment/packages")
async def get_payment_packages():
    """Get available payment packages."""
    return {
        "packages": PAYMENT_PACKAGES,
        "currency": "USD"
    }


@app.post("/payment/create-order")
async def create_payment_order(
    request: CreateOrderRequest,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Create Razorpay order for credit purchase."""
    
    # Validate package
    if request.package not in PAYMENT_PACKAGES:
        raise HTTPException(status_code=400, detail="Invalid package")
    
    pkg = PAYMENT_PACKAGES[request.package]
    user_id = current_user.get("user_id")
    
    # Convert to cents
    amount = int(pkg["price"] * 100)
    
    try:
        # ✅ FIXED: Shortened receipt to under 40 characters
        timestamp = int(datetime.utcnow().timestamp())
        # Take only last 8 chars of user_id if it's a UUID
        user_short = str(user_id)[-8:] if len(str(user_id)) > 8 else str(user_id)
        receipt = f"rcpt_{user_short}_{timestamp}"
        
        # Create Razorpay order
        order_data = {
            "amount": amount,
            "currency": pkg["currency"],
            "receipt": receipt,  # Now guaranteed under 40 chars
            "notes": {
                "user_id": str(user_id),
                "package": request.package,
                "credits": pkg["credits"]
            }
        }
        
        razorpay_order = razorpay_client.order.create(data=order_data)
        
        # Save to database
        purchase = CreditPurchase(
            user_id=user_id,
            razorpay_order_id=razorpay_order["id"],
            package=request.package,
            amount=pkg["price"],
            credits_purchased=pkg["credits"],
            payment_status="created",
            receipt=receipt  # Store the receipt
        )
        
        db.add(purchase)
        db.commit()
        db.refresh(purchase)
        
        return {
            "order_id": razorpay_order["id"],
            "amount": amount,
            "currency": pkg["currency"],
            "key_id": os.getenv("RAZORPAY_KEY_ID"),
            "package": request.package,
            "credits": pkg["credits"]
        }
        
    except Exception as e:
        print(f"Error creating order: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to create order: {str(e)}")
    
    
from razorpay.errors import SignatureVerificationError
import hmac
import hashlib

@app.post("/payment/verify")
async def verify_payment(
    request: VerifyPaymentRequest,  # ✅ CHANGED: Now accepts JSON body
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Verify payment and add credits."""
    
    user_id = current_user.get("user_id")
    
    # Verify signature
    try:
        params_dict = {
            'razorpay_order_id': request.razorpay_order_id,
            'razorpay_payment_id': request.razorpay_payment_id,
            'razorpay_signature': request.razorpay_signature
        }
        
        razorpay_client.utility.verify_payment_signature(params_dict)
        
    except SignatureVerificationError:
        raise HTTPException(status_code=400, detail="Invalid signature")
    
    # Find purchase
    purchase = db.query(CreditPurchase).filter(
        CreditPurchase.razorpay_order_id == request.razorpay_order_id,
        CreditPurchase.user_id == user_id
    ).first()
    
    if not purchase:
        raise HTTPException(status_code=404, detail="Purchase not found")
    
    # Check if already processed
    if purchase.payment_status == "paid":
        return {
            "success": True,
            "message": "Payment already processed",
            "credits_added": purchase.credits_purchased,
            "total_credits": db.query(User).filter(User.id == user_id).first().message_credits
        }
    
    # Update purchase
    purchase.razorpay_payment_id = request.razorpay_payment_id
    purchase.razorpay_signature = request.razorpay_signature
    purchase.payment_status = "paid"
    purchase.paid_at = datetime.utcnow()
    
    # Add credits to user
    user = db.query(User).filter(User.id == user_id).first()
    user.message_credits += purchase.credits_purchased
    user.total_credits_purchased += purchase.credits_purchased
    user.total_spent += purchase.amount
    user.is_premium = True
    user.last_purchase_date = datetime.utcnow()
    
    db.commit()
    
    return {
        "success": True,
        "message": "Payment successful!",
        "credits_added": purchase.credits_purchased,
        "total_credits": user.message_credits
    }


@app.get("/payment/packages")
async def get_payment_packages():
    """Get available payment packages."""
    return {
        "packages": PAYMENT_PACKAGES,
        "currency": "USD"
    }
    
@app.get("/user/credits")
async def get_user_credits(
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get user's credit balance."""
    
    user_id = current_user.get("user_id")
    user = db.query(User).filter(User.id == user_id).first()
    
    return {
        "is_premium": user.is_premium,
        "message_credits": user.message_credits,  # Changed key to match frontend
        "total_purchased": user.total_credits_purchased,
        "total_spent": float(user.total_spent),
        "last_purchase": user.last_purchase_date.isoformat() if user.last_purchase_date else None
    }


@app.get("/user/purchase-history")
async def get_purchase_history(
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get purchase history."""
    
    user_id = current_user.get("user_id")
    
    purchases = db.query(CreditPurchase).filter(
        CreditPurchase.user_id == user_id,
        CreditPurchase.payment_status == "paid"
    ).order_by(CreditPurchase.paid_at.desc()).all()
    
    return {
        "purchases": [
            {
                "id": p.id,
                "package": p.package,
                "credits": p.credits_purchased,
                "amount": float(p.amount),
                "currency": p.currency,
                "date": p.paid_at.isoformat(),
                "payment_id": p.razorpay_payment_id
            }
            for p in purchases
        ]
    }

# if __name__ == "__main__":
#     import uvicorn
#     port = int(os.getenv("PORT", 8082))
#     # uvicorn.run(app, host="0.0.0.0", port=port, log_level="info")
#     uvicorn.run(app, host="0.0.0.0", port=8082, ssl_keyfile="server.key",ssl_certfile="server.crt", log_level="info")

# import uvicorn
# uvicorn.run(app, host="0.0.0.0", port=8082, log_level="info")